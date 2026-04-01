import os
import re
import json
import time
import hashlib
import random
import threading
from io import BytesIO
from pathlib import Path
from typing import List, Tuple
from difflib import SequenceMatcher, ndiff
import datetime

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor

# Optional metrics deps
try:
    import sacrebleu
except Exception:
    sacrebleu = None

try:
    from bert_score import score as bertscore_score
except Exception:
    bertscore_score = None

# Optional plotting
try:
    import matplotlib.pyplot as plt  # noqa: F401
    HAVE_MPL = True
except Exception:
    HAVE_MPL = False

# Optional OpenAI
try:
    import openai
except Exception:
    openai = None


# ---------------- Proof-of-life banner ----------------
try:
    THIS_FILE = os.path.abspath(__file__)
    LAST_EDIT = datetime.datetime.fromtimestamp(os.path.getmtime(THIS_FILE))
except Exception:
    THIS_FILE = "interactive_session"
    LAST_EDIT = datetime.datetime.now()


# ---------------- Storage ----------------
DATA_DIR = Path("./data")
DATA_DIR.mkdir(exist_ok=True)

EXERCISES_FILE = DATA_DIR / "exercises.json"
SUBMISSIONS_FILE = DATA_DIR / "submissions.json"
LEADERBOARD_FILE = DATA_DIR / "leaderboard.json"
SUBMISSION_LOG_FILE = DATA_DIR / "submission_log.json"

LOC_STICKERS_FILE = DATA_DIR / "loc_stickers.json"
STICKER_IMG_DIR = DATA_DIR / "stickers"
STICKER_IMG_DIR.mkdir(exist_ok=True)

MQM_FILE = DATA_DIR / "mqm_assessments.json"
MQM_CONFIG_FILE = DATA_DIR / "mqm_config.json"

_lock = threading.Lock()


# ---------------- JSON helpers ----------------
def load_json(file: Path):
    file = Path(file)
    if file.exists():
        with file.open("r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return {}
    return {}


def save_json(file: Path, data):
    with _lock:
        tmp = Path(str(file) + ".tmp")
        with tmp.open("w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        tmp.replace(file)


def append_submission(student_name: str, ex_id: str, submission_data: dict):
    try:
        log_data = load_json(SUBMISSION_LOG_FILE)
        if not isinstance(log_data, list):
            log_data = []
        log_data.append({
            "student_name": student_name,
            "exercise_id": ex_id,
            "timestamp": datetime.datetime.now().isoformat(),
            "submission": submission_data,
        })
        save_json(SUBMISSION_LOG_FILE, log_data)
    except Exception:
        pass


# ---------------- Secrets / Auth ----------------
def get_secret(name, default=""):
    value = os.getenv(name, "")
    if value:
        return value
    try:
        if hasattr(st, "secrets") and name in st.secrets:
            return st.secrets[name]
    except Exception:
        pass
    return default


FALLBACK_PLAIN = "admin123"


def check_password(typed: str) -> bool:
    try:
        instructor_plain = get_secret("INSTRUCTOR_PASSWORD_PLAIN", "")
        instructor_sha256 = get_secret("INSTRUCTOR_PASSWORD_SHA256", "")
        instructor_dev_mode = get_secret("INSTRUCTOR_DEV_MODE", "0") == "1"

        if instructor_sha256:
            h = hashlib.sha256(typed.encode("utf-8")).hexdigest()
            return h == instructor_sha256

        if instructor_plain:
            return typed == instructor_plain

        if instructor_dev_mode:
            return typed == FALLBACK_PLAIN

        return False
    except Exception:
        return False


def has_instructor_auth_config() -> bool:
    return bool(
        get_secret("INSTRUCTOR_PASSWORD_PLAIN", "")
        or get_secret("INSTRUCTOR_PASSWORD_SHA256", "")
        or get_secret("INSTRUCTOR_DEV_MODE", "0") == "1"
    )


# ---------------- AI backend diagnostics ----------------
def get_ai_backend_status():
    openai_key = get_secret("OPENAI_API_KEY", "")
    openai_available = bool(openai_key)
    openai_model = get_secret("OPENAI_MODEL", "gpt-4.1-mini")

    hf_token = get_secret("HF_API_TOKEN", "")
    hf_available = bool(hf_token)

    return {
        "openai": openai_available,
        "openai_model": openai_model if openai_available else None,
        "hf": hf_available,
    }


# ---------------- Tokenization & Edit Helpers ----------------
TOKEN_RE = re.compile(r"\w+|[^\w\s]", re.UNICODE)


def _tokenize(s: str) -> List[str]:
    return TOKEN_RE.findall(s or "")


def compute_edit_details(mt_text: str, student_text: str) -> Tuple[int, int, int]:
    mt_tokens = _tokenize(mt_text)
    st_tokens = _tokenize(student_text)
    matcher = SequenceMatcher(None, mt_tokens, st_tokens)

    additions = deletions = replacements = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "insert":
            additions += (j2 - j1)
        elif tag == "delete":
            deletions += (i2 - i1)
        elif tag == "replace":
            replacements += max(i2 - i1, j2 - j1)

    total_edits = additions + deletions + replacements
    return additions, deletions, total_edits


# ---------------- Metrics ----------------
def evaluate_translation(student_text, mt_text=None, reference=None, task_type="Translate", source_text=""):
    src_len = max(1, len(_tokenize(source_text)))
    tgt_len = len(_tokenize(student_text))
    length_ratio = round(tgt_len / src_len, 3)

    if task_type == "Post-edit MT" and mt_text:
        additions, deletions, edits = compute_edit_details(mt_text, student_text)
    else:
        additions = deletions = edits = 0

    bleu = chrf = bert_f1 = None
    if reference:
        refs = [reference]
        try:
            if sacrebleu:
                bleu = float(sacrebleu.corpus_bleu([student_text], [refs]).score)
                chrf = float(sacrebleu.corpus_chrf([student_text], [refs]).score)
        except Exception:
            bleu = None
            chrf = None
        try:
            if bertscore_score:
                _, _, f1 = bertscore_score([student_text], [reference], lang="en")
                bert_f1 = float(f1.mean().item())
        except Exception:
            bert_f1 = None

    return {
        "length_ratio": length_ratio,
        "BLEU": None if bleu is None else round(bleu, 2),
        "chrF++": None if chrf is None else round(chrf, 2),
        "BERTScore_F1": None if bert_f1 is None else round(bert_f1, 3),
        "additions": additions,
        "deletions": deletions,
        "edits": edits,
    }


# ---------------- MQM ----------------
def load_mqm_config():
    cfg = load_json(MQM_CONFIG_FILE)
    if not cfg:
        cfg = {
            "categories": [
                "Accuracy",
                "Fluency",
                "Terminology",
                "Style",
                "Locale convention",
                "Design / markup",
                "Omission",
                "Addition",
                "Mistranslation",
                "Grammar",
                "Spelling / punctuation",
            ],
            "severity_weights": {
                "minor": 1,
                "major": 5,
                "critical": 10,
            },
        }
        save_json(MQM_CONFIG_FILE, cfg)
    return cfg


def load_mqm_assessments():
    data = load_json(MQM_FILE)
    return data if isinstance(data, dict) else {}


def save_mqm_assessment(student_name: str, ex_id: str, assessment: dict):
    data = load_mqm_assessments()
    if student_name not in data:
        data[student_name] = {}
    data[student_name][ex_id] = assessment
    save_json(MQM_FILE, data)


def get_mqm_assessment(student_name: str, ex_id: str):
    data = load_mqm_assessments()
    return data.get(student_name, {}).get(ex_id, {})


def compute_mqm_score(error_list, severity_weights):
    total_penalty = 0
    per_category = {}
    per_severity = {"minor": 0, "major": 0, "critical": 0}

    for err in error_list:
        sev = (err.get("severity") or "").strip().lower()
        cat = (err.get("category") or "").strip() or "Unspecified"
        penalty = severity_weights.get(sev, 0)

        total_penalty += penalty

        per_category.setdefault(cat, {"count": 0, "penalty": 0})
        per_category[cat]["count"] += 1
        per_category[cat]["penalty"] += penalty

        if sev in per_severity:
            per_severity[sev] += 1
        else:
            per_severity[sev] = per_severity.get(sev, 0) + 1

    score = max(0, 100 - total_penalty)
    return {
        "penalty": total_penalty,
        "score": score,
        "error_count": len(error_list),
        "per_category": per_category,
        "per_severity": per_severity,
    }


def mqm_errors_to_df(mqm_data: dict) -> pd.DataFrame:
    errors = mqm_data.get("errors", [])
    rows = []
    for i, e in enumerate(errors, start=1):
        rows.append({
            "#": i,
            "Category": e.get("category", ""),
            "Severity": e.get("severity", ""),
            "Span / Issue": e.get("span", ""),
            "Comment": e.get("comment", ""),
        })
    return pd.DataFrame(rows)


def mqm_category_summary_df(mqm_data: dict) -> pd.DataFrame:
    per_category = mqm_data.get("per_category", {})
    rows = []
    for cat, vals in per_category.items():
        rows.append({
            "Category": cat,
            "Error Count": vals.get("count", 0),
            "Penalty": vals.get("penalty", 0),
        })
    if not rows:
        return pd.DataFrame(columns=["Category", "Error Count", "Penalty"])
    return pd.DataFrame(rows).sort_values(["Penalty", "Error Count"], ascending=[False, False])


def render_mqm_summary(mqm_data: dict):
    if not mqm_data:
        st.info("No MQM assessment available yet.")
        return

    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("MQM Score", mqm_data.get("score", "—"))
    with c2:
        st.metric("MQM Penalty", mqm_data.get("penalty", "—"))
    with c3:
        st.metric("MQM Error Count", mqm_data.get("error_count", "—"))

    cat_df = mqm_category_summary_df(mqm_data)
    if not cat_df.empty:
        st.markdown("**MQM Category Summary**")
        st.dataframe(cat_df, use_container_width=True)

    err_df = mqm_errors_to_df(mqm_data)
    if not err_df.empty:
        st.markdown("**MQM Error Log**")
        st.dataframe(err_df, use_container_width=True)


def build_mqm_export_rows(mqm_all: dict):
    rows = []
    for student, student_data in mqm_all.items():
        for ex_id, mqm in student_data.items():
            errors = mqm.get("errors", [])
            if errors:
                for idx, err in enumerate(errors, start=1):
                    rows.append({
                        "Student": student,
                        "Exercise": ex_id,
                        "MQM Score": mqm.get("score"),
                        "MQM Penalty": mqm.get("penalty"),
                        "MQM Error Count": mqm.get("error_count"),
                        "Overall Comment": mqm.get("overall_comment", ""),
                        "Error #": idx,
                        "Category": err.get("category", ""),
                        "Severity": err.get("severity", ""),
                        "Span / Issue": err.get("span", ""),
                        "Comment": err.get("comment", ""),
                        "Assessed At": mqm.get("assessed_at", ""),
                    })
            else:
                rows.append({
                    "Student": student,
                    "Exercise": ex_id,
                    "MQM Score": mqm.get("score"),
                    "MQM Penalty": mqm.get("penalty"),
                    "MQM Error Count": mqm.get("error_count"),
                    "Overall Comment": mqm.get("overall_comment", ""),
                    "Error #": "",
                    "Category": "",
                    "Severity": "",
                    "Span / Issue": "",
                    "Comment": "",
                    "Assessed At": mqm.get("assessed_at", ""),
                })
    return rows


def build_mqm_overview_df(mqm_all: dict) -> pd.DataFrame:
    rows = []
    for student, student_data in mqm_all.items():
        for ex_id, mqm in student_data.items():
            rows.append({
                "Student": student,
                "Exercise": ex_id,
                "MQM Score": mqm.get("score"),
                "MQM Penalty": mqm.get("penalty"),
                "MQM Error Count": mqm.get("error_count"),
                "Minor Errors": mqm.get("per_severity", {}).get("minor", 0),
                "Major Errors": mqm.get("per_severity", {}).get("major", 0),
                "Critical Errors": mqm.get("per_severity", {}).get("critical", 0),
                "Assessed At": mqm.get("assessed_at", ""),
            })
    return pd.DataFrame(rows)


# ---------------- Track Changes ----------------
def _join_tokens_for_display(tokens: List[str]) -> str:
    out = " ".join(tokens)
    out = re.sub(r"\s+([.,!?;:])", r"\1", out)
    return out


def diff_text(baseline: str, student_text: str) -> str:
    differ = ndiff(_tokenize(baseline), _tokenize(student_text))
    parts = []
    for w in differ:
        token = w[2:]
        if w.startswith("- "):
            parts.append(f"<span style='color:#c00;text-decoration:line-through'>{token}</span>")
        elif w.startswith("+ "):
            parts.append(f"<span style='color:#080'>{token}</span>")
        else:
            parts.append(token)
    return _join_tokens_for_display(parts)


INVALID_XML_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def _safe_docx_text(value) -> str:
    if value is None:
        return ""
    if not isinstance(value, str):
        value = str(value)
    return INVALID_XML_RE.sub("", value)


def add_diff_to_doc(doc: Document, baseline: str, student_text: str):
    differ = ndiff(_tokenize(baseline), _tokenize(student_text))
    p = doc.add_paragraph()
    for w in differ:
        token = _safe_docx_text(w[2:])
        if w.startswith("- "):
            run = p.add_run(token + " ")
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif w.startswith("+ "):
            run = p.add_run(token + " ")
            run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            p.add_run(token + " ")


# ---------------- Exports ----------------
def export_student_word(submissions, student_name):
    doc = Document()
    doc.add_heading(_safe_docx_text(f"Student: {student_name}"), 0)
    subs = submissions.get(student_name, {})
    mqm_all = load_mqm_assessments()

    for ex_id, sub in subs.items():
        doc.add_heading(_safe_docx_text(f"Exercise {ex_id}"), level=1)
        doc.add_paragraph("Source Text:")
        doc.add_paragraph(_safe_docx_text(sub.get("source_text", "")))

        if sub.get("mt_text"):
            doc.add_paragraph("MT Output:")
            doc.add_paragraph(_safe_docx_text(sub.get("mt_text", "")))

        if sub.get("task_type") == "Post-edit MT":
            doc.add_paragraph("Student Submission (Track Changes):")
            base = sub.get("mt_text", "") or ""
            add_diff_to_doc(doc, base, sub.get("student_text", ""))
        else:
            doc.add_paragraph("Student Submission:")
            doc.add_paragraph(_safe_docx_text(sub.get("student_text", "")))

        metrics = sub.get("metrics", {})
        doc.add_paragraph(_safe_docx_text(f"Metrics: {metrics}"))
        doc.add_paragraph(_safe_docx_text(f"Task Type: {sub.get('task_type', '')}"))
        doc.add_paragraph(_safe_docx_text(f"Time Spent: {sub.get('time_spent_sec', 0):.2f} sec"))
        doc.add_paragraph(_safe_docx_text(f"Characters Typed: {sub.get('keystrokes', 0)}"))

        mqm = mqm_all.get(student_name, {}).get(ex_id, {})
        if mqm:
            doc.add_paragraph("MQM Assessment:")
            doc.add_paragraph(_safe_docx_text(f"MQM Score: {mqm.get('score', '')}"))
            doc.add_paragraph(_safe_docx_text(f"MQM Penalty: {mqm.get('penalty', '')}"))
            doc.add_paragraph(_safe_docx_text(f"MQM Error Count: {mqm.get('error_count', '')}"))

            if mqm.get("overall_comment"):
                doc.add_paragraph(_safe_docx_text(f"Overall Comment: {mqm.get('overall_comment', '')}"))

            per_category = mqm.get("per_category", {})
            if per_category:
                doc.add_paragraph("MQM Category Totals:")
                for cat, vals in per_category.items():
                    doc.add_paragraph(_safe_docx_text(
                        f"{cat}: count={vals.get('count', 0)}, penalty={vals.get('penalty', 0)}"
                    ))

            for err in mqm.get("errors", []):
                line = (
                    f"Category: {err.get('category', '')} | "
                    f"Severity: {err.get('severity', '')} | "
                    f"Span: {err.get('span', '')} | "
                    f"Comment: {err.get('comment', '')}"
                )
                doc.add_paragraph(_safe_docx_text(line))

        if sub.get("reflection"):
            doc.add_paragraph("Reflection:")
            doc.add_paragraph(_safe_docx_text(sub.get("reflection")))

        doc.add_paragraph("---")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_submissions_export_df(submissions):
    mqm_all = load_mqm_assessments()
    rows = []

    for student, subs in submissions.items():
        for ex_id, sub in subs.items():
            m = sub.get("metrics", {})
            mqm = mqm_all.get(student, {}).get(ex_id, {})

            rows.append({
                "Student": student,
                "Exercise": ex_id,
                "Task Type": sub.get("task_type", ""),
                "Source Text": sub.get("source_text", ""),
                "MT Output": sub.get("mt_text", ""),
                "Student Submission": sub.get("student_text", ""),
                "Reflection": sub.get("reflection", ""),
                "Length Ratio": m.get("length_ratio"),
                "BLEU": m.get("BLEU"),
                "chrF++": m.get("chrF++"),
                "BERTScore_F1": m.get("BERTScore_F1"),
                "Additions": m.get("additions"),
                "Deletions": m.get("deletions"),
                "Edits": m.get("edits"),
                "Time Spent (s)": sub.get("time_spent_sec", 0),
                "Characters Typed": sub.get("keystrokes", 0),
                "MQM Score": mqm.get("score"),
                "MQM Penalty": mqm.get("penalty"),
                "MQM Error Count": mqm.get("error_count"),
                "MQM Overall Comment": mqm.get("overall_comment", ""),
            })

    return pd.DataFrame(rows)


def build_leaderboard_df():
    leaderboard = load_json(LEADERBOARD_FILE)
    if not isinstance(leaderboard, dict):
        leaderboard = {}
    items = sorted(leaderboard.items(), key=lambda x: x[1], reverse=True)
    return pd.DataFrame(items, columns=["Student", "Points"])


def export_summary_excel(submissions):
    submissions_df = build_submissions_export_df(submissions)
    mqm_all = load_mqm_assessments()
    mqm_df = pd.DataFrame(build_mqm_export_rows(mqm_all))
    leaderboard_df = build_leaderboard_df()

    buf = BytesIO()
    try:
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            submissions_df.to_excel(writer, sheet_name="Submissions", index=False)
            mqm_df.to_excel(writer, sheet_name="MQM", index=False)
            leaderboard_df.to_excel(writer, sheet_name="Leaderboard", index=False)
    except Exception:
        with pd.ExcelWriter(buf) as writer:
            submissions_df.to_excel(writer, sheet_name="Submissions", index=False)
            mqm_df.to_excel(writer, sheet_name="MQM", index=False)
            leaderboard_df.to_excel(writer, sheet_name="Leaderboard", index=False)

    buf.seek(0)
    return buf


# ---------------- Gamification ----------------
def load_leaderboard():
    data = load_json(LEADERBOARD_FILE)
    return data if isinstance(data, dict) else {}


def update_leaderboard(student_name, points):
    leaderboard = load_leaderboard()
    leaderboard[student_name] = leaderboard.get(student_name, 0) + points
    save_json(LEADERBOARD_FILE, leaderboard)


def show_leaderboard():
    leaderboard = load_leaderboard()
    st.subheader("Leaderboard")
    if leaderboard:
        items = sorted(leaderboard.items(), key=lambda x: x[1], reverse=True)
        df = pd.DataFrame(items, columns=["Student", "Points"])
        st.dataframe(df, use_container_width=True)
    else:
        st.info("No leaderboard data yet.")


# ---------------- Instructor sticker manager ----------------
def localisation_sticker_manager():
    st.subheader("Instructor – Sticker / text / image localisation tasks")

    pwd = st.text_input("Instructor password", type="password", key="loc_sticker_pwd")
    if not check_password(pwd):
        if not has_instructor_auth_config():
            st.warning(
                "Instructor password is not configured. Set INSTRUCTOR_PASSWORD_PLAIN or "
                "INSTRUCTOR_PASSWORD_SHA256 (or INSTRUCTOR_DEV_MODE=1 for local testing)."
            )
        else:
            st.info("Enter instructor password to manage these tasks.")
        return

    loc_stickers = load_json(LOC_STICKERS_FILE)
    if not isinstance(loc_stickers, dict):
        loc_stickers = {}

    sticker_ids = ["New task"] + sorted(loc_stickers.keys())
    selection = st.selectbox("Choose task", sticker_ids, key="loc_sticker_select")

    if selection != "New task" and selection in loc_stickers:
        current = loc_stickers[selection]
        default_title = current.get("title", "")
        default_instr = current.get("instructions", "")
        default_text = current.get("content_text", "")
        default_url = current.get("image_url", "") if current.get("image_type") == "url" else ""

        st.markdown("**Current preview for students:**")
        if default_text:
            st.markdown("**Text to localise:**")
            st.write(default_text)
        if current.get("image_type") == "uploaded":
            img_path = current.get("image_path", "")
            if img_path and Path(img_path).exists():
                st.image(str(img_path))
            else:
                st.warning("Image file not found on server.")
        elif current.get("image_type") == "url":
            st.image(current.get("image_url", ""))
    else:
        default_title = ""
        default_instr = ""
        default_text = ""
        default_url = ""

    with st.form("loc_sticker_form"):
        title = st.text_input("Task title", value=default_title)
        content_text = st.text_area("Text to be localised (optional)", value=default_text, height=120)
        instructions = st.text_area(
            "Instructions for students (what to do with this text / image)",
            value=default_instr,
            height=120,
        )

        st.write("Sticker / image (choose either URL or upload, both optional):")
        image_url = st.text_input("Image URL (optional)", value=default_url)
        uploaded = st.file_uploader("Or upload an image file", type=["png", "jpg", "jpeg", "webp"])

        col1, col2 = st.columns(2)
        with col1:
            save_btn = st.form_submit_button("Save / Update task")
        with col2:
            delete_btn = st.form_submit_button("Delete this task")

    loc_stickers = load_json(LOC_STICKERS_FILE)
    if not isinstance(loc_stickers, dict):
        loc_stickers = {}

    if save_btn:
        if not title.strip() and not content_text.strip() and not image_url and not uploaded:
            st.warning("Please provide at least a title, some text, or an image before saving.")
            return

        if selection != "New task" and selection in loc_stickers:
            task_id = selection
        else:
            existing_nums = []
            for sid in loc_stickers.keys():
                m = re.match(r"STK_(\d+)$", sid)
                if m:
                    existing_nums.append(int(m.group(1)))
            next_num = max(existing_nums + [0]) + 1
            task_id = f"STK_{next_num:03d}"

        image_type = None
        image_path = ""
        image_url_final = ""

        if uploaded is not None:
            safe_name = re.sub(r"[^\w\.\-]", "_", uploaded.name)
            file_path = STICKER_IMG_DIR / f"{task_id}_{safe_name}"
            with open(file_path, "wb") as f:
                f.write(uploaded.getbuffer())
            image_type = "uploaded"
            image_path = str(file_path)
        elif image_url:
            image_type = "url"
            image_url_final = image_url

        if task_id in loc_stickers and image_type is None:
            image_type = loc_stickers[task_id].get("image_type")
            image_path = loc_stickers[task_id].get("image_path", "")
            image_url_final = loc_stickers[task_id].get("image_url", "")

        loc_stickers[task_id] = {
            "title": title.strip(),
            "instructions": instructions.strip(),
            "content_text": content_text.strip(),
            "image_type": image_type,
            "image_path": image_path,
            "image_url": image_url_final,
            "created_at": datetime.datetime.now().isoformat(),
        }
        save_json(LOC_STICKERS_FILE, loc_stickers)
        st.success(f"Task {task_id} saved.")

    if delete_btn and selection != "New task" and selection in loc_stickers:
        loc_stickers.pop(selection, None)
        save_json(LOC_STICKERS_FILE, loc_stickers)
        st.success(f"Task {selection} deleted.")


# ---------------- Optional AI exercise generator ----------------
def ai_generate_text(prompt):
    hf_token = get_secret("HF_API_TOKEN", "")
    if not hf_token:
        return None

    try:
        import requests
        headers = {"Authorization": f"Bearer {hf_token}"}
        payload = {"inputs": prompt}
        response = requests.post(
            "https://api-inference.huggingface.co/models/gpt2",
            headers=headers,
            json=payload,
            timeout=15,
        )
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, list) and data and "generated_text" in data[0]:
                return data[0]["generated_text"]
    except Exception:
        pass
    return None


# ---------------- Evidence-based Linguistic Hints ----------------
AR_LETTERS = r"\u0600-\u06FF"


def _tokenize_words(text: str):
    return re.findall(
        r"[A-Za-z" + AR_LETTERS + r"]+[’'\-]?[A-Za-z" + AR_LETTERS + r"]+|\d+(?:[.,]\d+)?",
        text,
    )


def _likely_terms(source_text: str):
    terms = set()

    for q in re.findall(r"[\"“”‘’'`«»](.+?)[\"“”‘’'`«»]", source_text):
        for w in _tokenize_words(q):
            if len(w) >= 3:
                terms.add(w)

    for w in _tokenize_words(source_text):
        if re.match(r"[A-Z][A-Za-z\-]+$", w):
            terms.add(w)
        elif re.match(r"[A-Z0-9\-]{3,}$", w):
            terms.add(w)
        elif "-" in w or re.search(r"\d", w):
            terms.add(w)
        elif re.match(r"[" + AR_LETTERS + r"]{4,}$", w):
            terms.add(w)

    return terms


def _short_list(items, n=4):
    items = list(items)
    if not items:
        return ""
    if len(items) <= n:
        return " | ".join(items)
    return " | ".join(items[:n]) + f" … (+{len(items)-n} more)"


def quick_linguistic_hints(source_text: str, student_text: str):
    hints = []
    try:
        src_nums = set(re.findall(r"\d+(?:[.,]\d+)?", source_text))
        tgt_nums = set(re.findall(r"\d+(?:[.,]\d+)?", student_text))
        missing_nums = sorted(src_nums - tgt_nums, key=lambda x: (len(x), x))
        if missing_nums:
            hints.append({
                "rule": "numbers_missing",
                "message": "Some figures from the source didn’t appear in your text.",
                "evidence": f"Missing: {_short_list(missing_nums)}",
            })

        for sym_open, sym_close, label in [
            ("(", ")", "parentheses"),
            ("[", "]", "brackets"),
            ("{", "}", "braces"),
        ]:
            if (
                source_text.count(sym_open) != student_text.count(sym_open)
                or source_text.count(sym_close) != student_text.count(sym_close)
            ):
                hints.append({
                    "rule": f"{label}_unbalanced",
                    "message": f"{label.capitalize()} look unbalanced.",
                    "evidence": (
                        f"Source {sym_open}/{sym_close}: "
                        f"{source_text.count(sym_open)}/{source_text.count(sym_close)}; "
                        f"Your text: {student_text.count(sym_open)}/{student_text.count(sym_close)}"
                    ),
                })

        if source_text.count('"') != student_text.count('"'):
            hints.append({
                "rule": "quotes_unbalanced",
                "message": "Quotation marks may be unbalanced.",
                "evidence": f"Source quotes: {source_text.count(chr(34))}; Yours: {student_text.count(chr(34))}",
            })

        src_terms = _likely_terms(source_text)
        tgt_tokens = set(_tokenize_words(student_text))
        missing_terms = sorted([t for t in src_terms if t not in tgt_tokens], key=lambda x: (-len(x), x))
        if missing_terms:
            hints.append({
                "rule": "terms_missing",
                "message": "Some key terms/names from the source weren’t reflected.",
                "evidence": f"Examples: {_short_list(missing_terms)}",
            })
    except Exception:
        pass

    return hints


# ---------------- Adaptive Feedback ----------------
def generate_feedback(metrics: dict, task_type: str, source_text: str, student_text: str, extra_hints=None):
    msgs = []
    lr = metrics.get("length_ratio")
    edits = int(metrics.get("edits", 0) or 0)
    adds = int(metrics.get("additions", 0) or 0)
    dels = int(metrics.get("deletions", 0) or 0)
    bleu = metrics.get("BLEU")
    chrf = metrics.get("chrF++")

    if task_type == "Post-edit MT":
        if edits == 0:
            msgs.append((
                "edits_none",
                "No edits were applied to the MT output.",
                "Review the MT carefully—critical errors may remain.",
            ))
        elif edits > 20:
            msgs.append((
                "edits_many",
                f"High edit volume detected: {edits} edits (additions {adds}, deletions {dels}).",
                "Prioritize adequacy/accuracy first; avoid cosmetic rephrasing that doesn’t fix meaning.",
            ))

    if lr is not None:
        if lr < 0.80:
            msgs.append((
                "len_low",
                f"Length ratio is {lr:.2f} (target ~0.90–1.20).",
                "Your translation may be over-compressed—recheck for omitted content.",
            ))
        elif lr > 1.30:
            msgs.append((
                "len_high",
                f"Length ratio is {lr:.2f} (target ~0.90–1.20).",
                "Consider concision—trim redundancy and literal padding.",
            ))

    if bleu is not None and chrf is not None:
        if bleu < 30 <= chrf:
            msgs.append((
                "acc_low_flu_ok",
                f"chrF++ is {chrf:.1f} (fluency ok) but BLEU is {bleu:.1f} (accuracy lagging).",
                "Revisit terminology and key meaning units; cross-check against the source.",
            ))
        elif bleu >= 30 and chrf < 50:
            msgs.append((
                "flu_low_acc_ok",
                f"BLEU is {bleu:.1f} (accuracy acceptable) but chrF++ is {chrf:.1f} (fluency weak).",
                "Polish cohesion and flow—simplify long clauses and connectors.",
            ))
        elif bleu < 20:
            msgs.append((
                "both_low",
                f"BLEU is {bleu:.1f}.",
                "Start with adequacy: ensure all propositions are conveyed before stylistic edits.",
            ))

    if extra_hints:
        for h in extra_hints:
            rule = h.get("rule", "hint")
            msg = h.get("message", "")
            evd = h.get("evidence", "")
            msgs.append((rule, msg, evd))

    seen = set()
    final = []
    for key, text, detail in msgs:
        if key in seen:
            continue
        seen.add(key)
        if detail:
            final.append(f"• {text} — *{detail}*")
        else:
            final.append(f"• {text}")
        if len(final) >= 4:
            break

    return final


# ---------------- AI ----------------
def build_ai_feedback_prompt(source_text: str, mt_text: str, student_text: str, task_type: str) -> str:
    mt_block = mt_text if mt_text else "(no MT output – direct translation task)"
    return f"""
You are an expert English–Arabic translation trainer specialising in translation and MT post-editing.

TASK TYPE: {task_type}

SOURCE TEXT:
{source_text}

MT OUTPUT (if any):
{mt_block}

STUDENT VERSION:
{student_text}

Give concise feedback suitable for a university translation classroom. Please:
1) Comment on accuracy (meaning transfer).
2) Comment on register and appropriateness for the context.
3) Comment on idiomatic and culturally appropriate choices.
4) Highlight one or two concrete examples where the student could improve.
5) If useful, propose a short improved version of one or two sentences, not the entire text.

You may answer partly in Arabic where it helps the student, but keep the structure clear and concise.
"""


def ask_ai_tutor(system_prompt: str, user_prompt: str):
    openai_key = get_secret("OPENAI_API_KEY", "")
    openai_model = get_secret("OPENAI_MODEL", "gpt-4.1-mini")

    if openai and openai_key:
        try:
            if hasattr(openai, "OpenAI"):
                client = openai.OpenAI(api_key=openai_key)

                try:
                    resp = client.responses.create(
                        model=openai_model,
                        instructions=system_prompt,
                        input=user_prompt,
                        max_output_tokens=900,
                    )
                    text = getattr(resp, "output_text", None)
                    if text:
                        return text.strip()
                except Exception:
                    resp = client.chat.completions.create(
                        model=openai_model,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                        max_tokens=900,
                        temperature=0.3,
                    )
                    text = resp.choices[0].message.content
                    if text:
                        return text.strip()
            else:
                openai.api_key = openai_key
                resp = openai.ChatCompletion.create(
                    model=openai_model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    max_tokens=900,
                    temperature=0.3,
                )
                text = resp.choices[0].message["content"]
                if text:
                    return text.strip()
        except Exception as e:
            return f"OpenAI error: {e}"

    hf_token = get_secret("HF_API_TOKEN", "")
    if hf_token:
        try:
            import requests
            headers = {"Authorization": f"Bearer {hf_token}"}
            payload = {"inputs": system_prompt + "\n\n" + user_prompt}
            response = requests.post(
                "https://api-inference.huggingface.co/models/gpt2",
                headers=headers,
                json=payload,
                timeout=30,
            )
            if response.status_code == 200:
                data = response.json()
                if isinstance(data, list) and data and "generated_text" in data[0]:
                    return data[0]["generated_text"].strip()
            return f"Hugging Face error: status {response.status_code}"
        except Exception as e:
            return f"Hugging Face error: {e}"

    return None


def generate_ai_feedback(prompt: str):
    system_prompt = "You are a helpful, expert translation instructor."
    return ask_ai_tutor(system_prompt, prompt)


def build_ai_mqm_prompt(source_text: str, mt_text: str, student_text: str, categories: list) -> str:
    categories_text = ", ".join(categories)
    mt_block = mt_text if mt_text else "(none)"
    return f"""
You are an expert translation quality assessor using MQM principles.

Allowed MQM categories:
{categories_text}

Severity must be one of:
minor, major, critical

Task:
Read the source text, optional MT output, and the student submission.
Identify up to 8 likely MQM issues.
Return ONLY valid JSON as a list of objects.
Each object must have exactly these keys:
category, severity, span, comment

Rules:
- category must be one of the allowed categories above
- severity must be one of: minor, major, critical
- span should be a short problematic segment or issue label
- comment should be concise and specific
- do not include markdown
- do not include explanation outside JSON
- if there are no clear issues, return []

SOURCE TEXT:
{source_text}

MT OUTPUT:
{mt_block}

STUDENT SUBMISSION:
{student_text}
"""


def parse_ai_mqm_suggestions(ai_text: str, allowed_categories: list):
    if not ai_text:
        return []

    text = ai_text.strip()

    # Try direct JSON
    candidates = [text]

    # Try extracting array block
    match = re.search(r"(\[\s*{.*}\s*\])", text, re.DOTALL)
    if match:
        candidates.insert(0, match.group(1))

    for candidate in candidates:
        try:
            data = json.loads(candidate)
            if not isinstance(data, list):
                continue

            clean = []
            for item in data:
                if not isinstance(item, dict):
                    continue
                category = str(item.get("category", "")).strip()
                severity = str(item.get("severity", "")).strip().lower()
                span = str(item.get("span", "")).strip()
                comment = str(item.get("comment", "")).strip()

                if category not in allowed_categories:
                    continue
                if severity not in {"minor", "major", "critical"}:
                    continue
                if not span and not comment:
                    continue

                clean.append({
                    "category": category,
                    "severity": severity,
                    "span": span,
                    "comment": comment,
                })
            return clean
        except Exception:
            continue

    return []


# ---------------- Instructor MQM panel ----------------
def instructor_mqm_panel(exercises, submissions):
    st.subheader("MQM Assessment")

    cfg = load_mqm_config()
    categories = cfg.get("categories", [])
    severity_weights = cfg.get("severity_weights", {"minor": 1, "major": 5, "critical": 10})

    with st.expander("MQM Configuration", expanded=False):
        st.markdown("**Categories**")
        cat_text = st.text_area(
            "One category per line",
            value="\n".join(categories),
            height=180,
            key="mqm_cat_text",
        )

        st.markdown("**Severity weights**")
        c1, c2, c3 = st.columns(3)
        with c1:
            minor_w = st.number_input("Minor", min_value=0, value=int(severity_weights.get("minor", 1)), key="mqm_minor")
        with c2:
            major_w = st.number_input("Major", min_value=0, value=int(severity_weights.get("major", 5)), key="mqm_major")
        with c3:
            critical_w = st.number_input("Critical", min_value=0, value=int(severity_weights.get("critical", 10)), key="mqm_critical")

        if st.button("Save MQM configuration", key="save_mqm_cfg"):
            new_cfg = {
                "categories": [x.strip() for x in cat_text.splitlines() if x.strip()],
                "severity_weights": {
                    "minor": int(minor_w),
                    "major": int(major_w),
                    "critical": int(critical_w),
                },
            }
            save_json(MQM_CONFIG_FILE, new_cfg)
            st.success("MQM configuration saved.")
            cfg = new_cfg
            categories = cfg["categories"]
            severity_weights = cfg["severity_weights"]

    if not submissions:
        st.info("No student submissions yet for MQM assessment.")
        return

    student_names = sorted(submissions.keys())
    selected_student = st.selectbox("Student for MQM", student_names, key="mqm_student")

    student_subs = submissions.get(selected_student, {})
    if not student_subs:
        st.info("This student has no submissions.")
        return

    selected_ex = st.selectbox("Exercise for MQM", sorted(student_subs.keys()), key="mqm_ex")
    sub = student_subs[selected_ex]

    st.markdown("### Submission under review")
    st.markdown("**Source text:**")
    st.write(sub.get("source_text", ""))

    if sub.get("mt_text"):
        st.markdown("**MT output:**")
        st.write(sub.get("mt_text", ""))

    st.markdown("**Student submission:**")
    st.write(sub.get("student_text", ""))

    ai_key = f"mqm_ai_suggestions::{selected_student}::{selected_ex}"

    c_ai1, c_ai2 = st.columns([1, 2])
    with c_ai1:
        if st.button("Generate AI MQM Suggestions", key="gen_ai_mqm"):
            prompt = build_ai_mqm_prompt(
                sub.get("source_text", ""),
                sub.get("mt_text", "") or "",
                sub.get("student_text", ""),
                categories,
            )
            system_prompt = "You are a precise MQM assessor. Return only valid JSON."
            with st.spinner("Generating AI MQM suggestions..."):
                ai_text = ask_ai_tutor(system_prompt, prompt)
            suggestions = parse_ai_mqm_suggestions(ai_text, categories)
            st.session_state[ai_key] = suggestions
            if suggestions:
                st.success(f"Loaded {len(suggestions)} AI suggestion(s).")
            else:
                st.warning("No usable AI suggestions were returned.")
    with c_ai2:
        if ai_key in st.session_state and st.session_state[ai_key]:
            st.caption("AI suggestions were loaded. You can edit them below before saving.")

    existing = get_mqm_assessment(selected_student, selected_ex)
    existing_errors = existing.get("errors", [])
    ai_errors = st.session_state.get(ai_key, [])

    base_errors = ai_errors if ai_errors else existing_errors
    default_rows = max(3, len(base_errors) if base_errors else 3)

    row_count = st.number_input(
        "Number of MQM error rows",
        min_value=1,
        max_value=20,
        value=default_rows,
        step=1,
        key="mqm_rows",
    )

    errors = []
    for i in range(int(row_count)):
        st.markdown(f"**MQM Error {i + 1}**")
        old = base_errors[i] if i < len(base_errors) else {}

        col1, col2 = st.columns(2)
        with col1:
            options = categories if categories else ["Accuracy"]
            idx = options.index(old["category"]) if old.get("category") in options else 0
            category = st.selectbox(
                f"Category {i + 1}",
                options,
                index=idx,
                key=f"mqm_cat_{i}",
            )
        with col2:
            sev_options = ["minor", "major", "critical"]
            sev_idx = sev_options.index(old["severity"]) if old.get("severity") in sev_options else 0
            severity = st.selectbox(
                f"Severity {i + 1}",
                sev_options,
                index=sev_idx,
                key=f"mqm_sev_{i}",
            )

        span = st.text_input(f"Span / issue {i + 1}", value=old.get("span", ""), key=f"mqm_span_{i}")
        comment = st.text_area(f"Comment {i + 1}", value=old.get("comment", ""), height=70, key=f"mqm_comment_{i}")

        if span.strip() or comment.strip():
            errors.append({
                "category": category,
                "severity": severity,
                "span": span.strip(),
                "comment": comment.strip(),
            })

    overall_comment = st.text_area(
        "Overall MQM comment",
        value=existing.get("overall_comment", ""),
        height=120,
        key="mqm_overall_comment",
    )

    save_col, clear_col = st.columns(2)
    with save_col:
        if st.button("Save MQM assessment", key="save_mqm_assessment"):
            summary = compute_mqm_score(errors, severity_weights)
            payload = {
                "errors": errors,
                "overall_comment": overall_comment.strip(),
                "score": summary["score"],
                "penalty": summary["penalty"],
                "error_count": summary["error_count"],
                "per_category": summary["per_category"],
                "per_severity": summary["per_severity"],
                "severity_weights": severity_weights,
                "assessed_at": datetime.datetime.now().isoformat(),
            }
            save_mqm_assessment(selected_student, selected_ex, payload)
            st.success("MQM assessment saved.")
    with clear_col:
        if st.button("Clear AI suggestions", key="clear_ai_mqm"):
            st.session_state.pop(ai_key, None)
            st.success("AI suggestions cleared.")

    latest = get_mqm_assessment(selected_student, selected_ex)
    if latest:
        st.markdown("### Current MQM result")
        render_mqm_summary(latest)
        if latest.get("overall_comment"):
            st.markdown("**Overall comment:**")
            st.write(latest.get("overall_comment"))


def instructor_mqm_analytics_panel():
    st.subheader("MQM Analytics")
    mqm_all = load_mqm_assessments()
    overview_df = build_mqm_overview_df(mqm_all)

    if overview_df.empty:
        st.info("No MQM analytics available yet.")
        return

    st.markdown("**MQM Overview by Submission**")
    st.dataframe(overview_df, use_container_width=True)

    if "MQM Score" in overview_df.columns:
        try:
            chart_df = overview_df.copy()
            chart_df["Label"] = chart_df["Student"].astype(str) + " | " + chart_df["Exercise"].astype(str)
            st.markdown("**MQM Score by Submission**")
            st.bar_chart(chart_df.set_index("Label")[["MQM Score"]])
        except Exception:
            pass

    cat_rows = []
    for student, student_data in mqm_all.items():
        for ex_id, mqm in student_data.items():
            for cat, vals in mqm.get("per_category", {}).items():
                cat_rows.append({
                    "Student": student,
                    "Exercise": ex_id,
                    "Category": cat,
                    "Error Count": vals.get("count", 0),
                    "Penalty": vals.get("penalty", 0),
                })

    if cat_rows:
        cat_df = pd.DataFrame(cat_rows)
        summary_cat = cat_df.groupby("Category", as_index=False)[["Error Count", "Penalty"]].sum()
        summary_cat = summary_cat.sort_values(["Penalty", "Error Count"], ascending=[False, False])

        st.markdown("**MQM Totals by Category**")
        st.dataframe(summary_cat, use_container_width=True)

        try:
            st.markdown("**Category Penalty Chart**")
            st.bar_chart(summary_cat.set_index("Category")[["Penalty"]])
        except Exception:
            pass

        try:
            st.markdown("**Category Error Count Chart**")
            st.bar_chart(summary_cat.set_index("Category")[["Error Count"]])
        except Exception:
            pass


# ---------------- Instructor Dashboard ----------------
def instructor_dashboard():
    st.title("Instructor Dashboard")
    password = st.text_input("Enter instructor password", type="password")

    if not check_password(password):
        if not has_instructor_auth_config():
            st.warning(
                "Instructor password is not configured. "
                "Set INSTRUCTOR_PASSWORD_PLAIN or INSTRUCTOR_PASSWORD_SHA256 "
                "(or INSTRUCTOR_DEV_MODE=1 for local testing)."
            )
        else:
            st.warning("Incorrect password. Access denied.")
        return

    st.subheader("AI Feedback Diagnostics")
    status = get_ai_backend_status()

    openai_badge = "Not configured"
    hf_badge = "Not configured"

    if status["openai"]:
        openai_badge = f"Available (model: `{status['openai_model']}`)"
    if status["hf"]:
        hf_badge = "Available (HF_API_TOKEN set)"

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("**OpenAI (ChatGPT API)**")
        st.write(openai_badge)
    with col_b:
        st.markdown("**Hugging Face Inference API**")
        st.write(hf_badge)

    if not status["openai"] and not status["hf"]:
        st.info(
            "No AI backend is currently configured.\n\n"
            "- To use ChatGPT-based feedback, set `OPENAI_API_KEY` (and optional `OPENAI_MODEL`).\n"
            "- To use a Hugging Face model, set `HF_API_TOKEN`."
        )
    else:
        st.success("AI feedback is at least partially configured.")

    st.markdown("---")

    exercises = load_json(EXERCISES_FILE)
    if not isinstance(exercises, dict):
        exercises = {}

    submissions = load_json(SUBMISSIONS_FILE)
    if not isinstance(submissions, dict):
        submissions = {}

    st.subheader("Create / Edit / Delete Exercise")
    ex_ids = ["New"] + list(exercises.keys())
    selected_ex = st.selectbox("Select Exercise", ex_ids)

    if selected_ex != "New" and selected_ex in exercises:
        default_source = exercises[selected_ex].get("source_text", "")
        default_mt = exercises[selected_ex].get("mt_text", "") or ""
    else:
        default_source = ""
        default_mt = ""

    with st.form("exercise_form"):
        st_text = st.text_area("Source Text", value=default_source, height=150)
        mt_text = st.text_area("MT Output (optional)", value=default_mt, height=150)
        c1, c2, c3 = st.columns(3)
        with c1:
            save_btn = st.form_submit_button("Save Exercise")
        with c2:
            delete_btn = st.form_submit_button("Delete Exercise")
        with c3:
            gen_btn = st.form_submit_button("Generate AI Exercise")

    if save_btn:
        try:
            next_id = (
                str(max([int(k) for k in exercises.keys()] + [0]) + 1).zfill(3)
                if selected_ex == "New" else selected_ex
            )
        except Exception:
            next_id = "001" if selected_ex == "New" else selected_ex

        exercises[next_id] = {
            "source_text": st_text,
            "mt_text": (mt_text.strip() or None),
        }
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise saved. ID: {next_id}")

    if delete_btn and selected_ex != "New":
        exercises.pop(selected_ex, None)
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise {selected_ex} deleted.")

    if gen_btn:
        prompt = "Write a short culturally rich text for translation students."
        ai_text = ai_generate_text(prompt)
        new_text = ai_text if ai_text else f"This is AI-generated exercise {random.randint(1, 1000)}."
        new_mt = f"MT output for exercise {random.randint(1, 1000)}."
        try:
            next_id = str(max([int(k) for k in exercises.keys()] + [0]) + 1).zfill(3)
        except Exception:
            next_id = "001"
        exercises[next_id] = {"source_text": new_text, "mt_text": new_mt}
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise saved as ID {next_id}")

    st.subheader("Download Exercises")
    if exercises:
        for ex_id, ex in exercises.items():
            try:
                buf = BytesIO()
                doc = Document()
                doc.add_heading(f"Exercise {ex_id}", 0)
                doc.add_paragraph("Source Text:")
                doc.add_paragraph(ex.get("source_text", ""))
                if ex.get("mt_text"):
                    doc.add_paragraph("MT Output:")
                    doc.add_paragraph(ex.get("mt_text", ""))
                doc.save(buf)
                buf.seek(0)
                st.download_button(
                    f"Exercise {ex_id} (Word)",
                    buf,
                    file_name=f"Exercise_{ex_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception:
                st.info(f"Exercise {ex_id}: export not available (DOCX error).")
    else:
        st.info("No exercises yet.")

    st.subheader("Student Submissions & Exports")
    if submissions:
        student_choice = st.selectbox("Choose student", ["All"] + list(submissions.keys()))
        if student_choice != "All":
            try:
                buf = export_student_word(submissions, student_choice)
                safe_name = re.sub(r"[^\w\-]+", "_", student_choice)
                st.download_button(
                    f"Download {student_choice}'s Submissions (Word)",
                    buf,
                    file_name=f"{safe_name}_submissions.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"Word export failed: {e}")

        st.subheader("Download Summary Workbook (Excel)")
        try:
            excel_buf = export_summary_excel(submissions)
            st.download_button(
                "Download Excel Workbook",
                excel_buf,
                file_name="metrics_summary.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception as e:
            st.error(f"Excel export failed: {e}")

        try:
            st.subheader("Class Snapshot")
            rows = []
            mqm_all = load_mqm_assessments()

            for ex_id2 in exercises.keys():
                vals = []
                mqm_vals = []
                for student, subs in submissions.items():
                    sub = subs.get(ex_id2)
                    if sub:
                        m = sub.get("metrics", {})
                        if m.get("chrF++") is not None:
                            vals.append(m["chrF++"])
                        mqm = mqm_all.get(student, {}).get(ex_id2, {})
                        if mqm.get("score") is not None:
                            mqm_vals.append(mqm.get("score"))
                if vals or mqm_vals:
                    row = {"Exercise": ex_id2, "n": max(len(vals), len(mqm_vals))}
                    if vals:
                        row["chrF++ mean"] = round(sum(vals) / max(1, len(vals)), 2)
                    if mqm_vals:
                        row["MQM mean"] = round(sum(mqm_vals) / max(1, len(mqm_vals)), 2)
                    rows.append(row)

            if rows:
                st.dataframe(pd.DataFrame(rows), use_container_width=True)
            else:
                st.info("No metrics yet to summarize.")
        except Exception:
            st.info("Snapshot unavailable (aggregation error).")

        show_leaderboard()
    else:
        st.info("No submissions yet.")

    st.markdown("---")
    instructor_mqm_panel(exercises, submissions)
    st.markdown("---")
    instructor_mqm_analytics_panel()


# ---------------- Student Dashboard ----------------
def student_dashboard():
    st.title("Student Dashboard")

    exercises = load_json(EXERCISES_FILE)
    if not isinstance(exercises, dict):
        exercises = {}

    if not exercises:
        st.info("No exercises available yet. Please check back later.")
        return

    submissions = load_json(SUBMISSIONS_FILE)
    if not isinstance(submissions, dict):
        submissions = {}

    student_name = st.text_input("Enter your name")
    if not student_name:
        return

    if student_name not in submissions:
        submissions[student_name] = {}

    ex_id = st.selectbox("Choose Exercise", list(exercises.keys()))
    if not ex_id:
        return

    ex = exercises[ex_id]
    existing_sub = submissions.get(student_name, {}).get(ex_id, {})

    st.subheader("Source Text")
    st.markdown(
        f"<div style='font-family:Times New Roman;font-size:12pt;'>{ex.get('source_text', '')}</div>",
        unsafe_allow_html=True,
    )

    task_options = ["Translate"] if not ex.get("mt_text") else ["Translate", "Post-edit MT"]

    default_task_type = existing_sub.get("task_type", task_options[0])
    if default_task_type not in task_options:
        default_task_type = task_options[0]

    task_type = st.radio(
        "Task Type",
        task_options,
        horizontal=True,
        index=task_options.index(default_task_type),
    )

    if existing_sub.get("student_text"):
        initial_text = existing_sub.get("student_text", "")
    else:
        initial_text = "" if task_type == "Translate" else (ex.get("mt_text", "") or "")

    default_reflection = existing_sub.get("reflection", "")

    start_key = f"start_time_{student_name}_{ex_id}"
    keys_key = f"chars_{student_name}_{ex_id}"

    if start_key not in st.session_state:
        st.session_state[start_key] = time.time()
    if keys_key not in st.session_state:
        st.session_state[keys_key] = 0

    with st.form(key=f"exercise_form_{student_name}_{ex_id}"):
        student_text = st.text_area(
            "Type your translation / post-edit here",
            value=initial_text,
            height=300,
        )
        reflection = st.text_area(
            "Brief reflection (what changed / why?)",
            value=default_reflection,
            height=80,
        )
        submitted = st.form_submit_button("Submit")

    if submitted:
        time_spent = time.time() - st.session_state[start_key]
        st.session_state[keys_key] = len(student_text)

        metrics = evaluate_translation(
            student_text,
            mt_text=ex.get("mt_text"),
            reference=None,
            task_type=task_type,
            source_text=ex.get("source_text", ""),
        )

        submissions[student_name][ex_id] = {
            "source_text": ex.get("source_text", ""),
            "mt_text": ex.get("mt_text"),
            "student_text": student_text,
            "task_type": task_type,
            "time_spent_sec": round(time_spent, 2),
            "keystrokes": st.session_state[keys_key],
            "metrics": metrics,
            "reflection": reflection,
        }
        save_json(SUBMISSIONS_FILE, submissions)
        append_submission(student_name, ex_id, submissions[student_name][ex_id])

        points = 0
        try:
            if metrics.get("BLEU") is not None:
                points += int(metrics["BLEU"])
            if metrics.get("chrF++") is not None:
                points += int(metrics["chrF++"] / 2)
            if task_type == "Post-edit MT":
                points += max(0, 10 - int(metrics["edits"]))
        except Exception:
            pass

        update_leaderboard(student_name, points)
        st.success("Submission saved.")
        existing_sub = submissions[student_name][ex_id]

    if not existing_sub or not existing_sub.get("student_text", "").strip():
        st.info("Write your translation and click Submit first. Then AI feedback and AI Tutor will work.")
        return

    student_text = existing_sub.get("student_text", "")
    task_type = existing_sub.get("task_type", task_type)
    time_spent = existing_sub.get("time_spent_sec", 0)
    st.session_state[keys_key] = existing_sub.get("keystrokes", len(student_text))
    metrics = existing_sub.get("metrics", {})

    def _fmt(v):
        return "—" if v is None else v

    st.subheader("Your Metrics")
    st.markdown(f"""
- **Length Ratio** (target/src): {_fmt(metrics.get('length_ratio'))}
- **BLEU**: {_fmt(metrics.get('BLEU'))}
- **chrF++**: {_fmt(metrics.get('chrF++'))}
- **BERTScore F1**: {_fmt(metrics.get('BERTScore_F1'))}
- **Additions**: {_fmt(metrics.get('additions'))}
- **Deletions**: {_fmt(metrics.get('deletions'))}
- **Edits**: {_fmt(metrics.get('edits'))}
- **Time Spent**: {time_spent} sec
- **Characters Typed**: {st.session_state[keys_key]}
""")

    st.subheader("MQM Assessment")
    mqm = get_mqm_assessment(student_name, ex_id)
    if mqm:
        render_mqm_summary(mqm)
        if mqm.get("overall_comment"):
            st.markdown("**Instructor MQM Comment:**")
            st.write(mqm.get("overall_comment"))
    else:
        st.info("No MQM assessment has been added yet by the instructor.")

    extra = quick_linguistic_hints(ex.get("source_text", ""), student_text)
    feedback_msgs = generate_feedback(
        metrics,
        task_type,
        ex.get("source_text", ""),
        student_text,
        extra,
    )

    st.subheader("Adaptive Feedback (metrics-based)")
    if feedback_msgs:
        for m in feedback_msgs:
            st.markdown(m)
    else:
        st.info("No specific issues triggered. Focus on cohesion, clarity, and consistent terminology.")

    if task_type == "Post-edit MT":
        st.subheader("Track Changes")
        st.caption("Track changes: green = additions, red strike = deletions.")
        base = ex.get("mt_text", "") or ""
        st.markdown(diff_text(base, student_text), unsafe_allow_html=True)

    try:
        history = []
        for ex_id2, sub2 in submissions.get(student_name, {}).items():
            m2 = sub2.get("metrics", {})
            mqm2 = get_mqm_assessment(student_name, ex_id2)
            history.append({
                "Exercise": ex_id2,
                "BLEU": m2.get("BLEU"),
                "chrF++": m2.get("chrF++"),
                "Edits": m2.get("edits", 0),
                "MQM Score": mqm2.get("score"),
            })

        if history:
            st.subheader("Progress Overview")
            df_hist = pd.DataFrame(history)

            try:
                trend_cols = [c for c in ["BLEU", "chrF++", "MQM Score"] if c in df_hist.columns]
                if trend_cols:
                    st.line_chart(df_hist.set_index("Exercise")[trend_cols])
            except Exception:
                pass

            try:
                st.bar_chart(df_hist.set_index("Exercise")[["Edits"]])
            except Exception:
                pass
    except Exception:
        st.info("Progress charts unavailable.")

    show_leaderboard()

    st.subheader("Optional AI Feedback (experimental)")
    if st.button("Get AI feedback on your submission"):
        prompt = build_ai_feedback_prompt(
            ex.get("source_text", ""),
            ex.get("mt_text", "") or "",
            student_text,
            task_type,
        )
        with st.spinner("Requesting AI feedback..."):
            ai_text = generate_ai_feedback(prompt)

        if ai_text:
            st.markdown("### AI feedback / suggestion")
            st.write(ai_text)
        else:
            st.error("AI call returned no text. Check API key, model name, and OpenAI package version.")

    st.subheader("AI Tutor Chat")
    student_prompt = st.text_area(
        "Ask the AI tutor a question about your translation",
        height=120,
        placeholder="Example: Is my Arabic too literal? Suggest 2 more idiomatic options for the first sentence.",
        key=f"ai_tutor_prompt_{student_name}_{ex_id}",
    )

    if st.button("Ask AI Tutor"):
        if not student_prompt.strip():
            st.warning("Please enter a prompt.")
        else:
            system_prompt = """
You are an expert English-Arabic translation tutor for university students.
Help students improve their work.
Do not automatically complete the whole task unless explicitly requested.
Focus on accuracy, register, idiomaticity, and cultural appropriateness.
Keep responses clear and concise.
You may use both English and Arabic.
"""
            user_prompt = f"""
TASK TYPE: {task_type}

SOURCE TEXT:
{ex.get('source_text', '')}

MT OUTPUT:
{ex.get('mt_text', '') or '(none)'}

STUDENT SUBMISSION:
{student_text}

STUDENT QUESTION:
{student_prompt}
"""
            with st.spinner("Requesting AI response..."):
                ai_text = ask_ai_tutor(system_prompt, user_prompt)

            if ai_text:
                st.markdown("### AI Tutor Response")
                st.write(ai_text)
            else:
                st.error("AI call returned no text. Check API key, model name, and OpenAI package version.")


# ---------------- Localisation Lab ----------------
def localisation_lab():
    st.title("Localisation Lab")
    st.write(
        "Interactive exercises on localisation (English ↔ Arabic). "
        "Work here is saved to the same JSON/leaderboard as the core lab."
    )

    mode = st.sidebar.radio(
        "Localisation mode",
        ["Student view", "Instructor (manage sticker/text/image tasks)"],
        index=0,
        key="loc_mode",
    )

    if mode == "Instructor (manage sticker/text/image tasks)":
        localisation_sticker_manager()
        return

    student_name = st.text_input("Enter your name (for saving localisation work)")
    if not student_name:
        st.info("Please enter your name to start.")
        return

    submissions = load_json(SUBMISSIONS_FILE)
    if not isinstance(submissions, dict):
        submissions = {}

    if student_name not in submissions:
        submissions[student_name] = {}

    exercise = st.sidebar.selectbox(
        "Choose a localisation exercise",
        [
            "1 Translation vs Localisation",
            "2 Cultural Adaptation in Advertising",
            "3 Conventions: Dates, Units, Currency",
            "4 Tone & Website/App UX",
            "5 Post-editing: Error Detection",
            "6 App Store Description",
            "7 Strategy & Theory Reflection",
            "Sticker / text / image task (from instructor)",
        ],
        key="loc_ex_select",
    )

    def save_loc_submission(ex_id: str, source_text: str, main_text: str, reflection_text: str):
        if not main_text.strip():
            st.warning("Nothing to save yet — please write your main answer first.")
            return

        start_key = f"loc_start_{student_name}_{ex_id}"
        if start_key not in st.session_state:
            st.session_state[start_key] = time.time()

        time_spent = time.time() - st.session_state[start_key]
        keystrokes = len(main_text)

        metrics = evaluate_translation(
            main_text,
            mt_text=None,
            reference=None,
            task_type="Localisation",
            source_text=source_text,
        )

        submissions[student_name][ex_id] = {
            "source_text": source_text,
            "mt_text": None,
            "student_text": main_text,
            "task_type": "Localisation",
            "time_spent_sec": round(time_spent, 2),
            "keystrokes": keystrokes,
            "metrics": metrics,
            "reflection": reflection_text,
        }
        save_json(SUBMISSIONS_FILE, submissions)

        points = 15
        lr = metrics.get("length_ratio")
        try:
            if lr is not None and 0.8 <= lr <= 1.3:
                points += 5
        except Exception:
            pass

        update_leaderboard(student_name, points)
        st.success("Localisation submission saved and leaderboard updated.")

    if exercise == "1 Translation vs Localisation":
        st.subheader("Translation vs Localisation")
        source_text = "Welcome to our app! Enjoy lightning-fast delivery and unbeatable deals every Friday."
        st.write("Decide how you would localise this text for Arabic users.")
        main_text = st.text_area("Your localised version", key="loc1_main", height=180)
        reflection = st.text_area("Reflection", key="loc1_reflect", height=100)
        if st.button("Save this task", key="loc1_save"):
            save_loc_submission("LOC_1", source_text, main_text, reflection)

    elif exercise == "2 Cultural Adaptation in Advertising":
        st.subheader("Cultural Adaptation in Advertising")
        source_text = "Grab your Halloween special now and trick-or-treat yourself with 50% off!"
        st.write("Adapt this for an Arabic-speaking audience while keeping the persuasive purpose.")
        main_text = st.text_area("Your adapted version", key="loc2_main", height=180)
        reflection = st.text_area("Reflection", key="loc2_reflect", height=100)
        if st.button("Save this task", key="loc2_save"):
            save_loc_submission("LOC_2", source_text, main_text, reflection)

    elif exercise == "3 Conventions: Dates, Units, Currency":
        st.subheader("Conventions: Dates, Units, Currency")
        source_text = "Offer valid until 12/31/2025. Free shipping on orders above $75. Package weight: 3.5 lbs."
        st.write("Rewrite/localise the text using conventions suitable for Arabic readers.")
        main_text = st.text_area("Your localised version", key="loc3_main", height=180)
        reflection = st.text_area("Reflection", key="loc3_reflect", height=100)
        if st.button("Save this task", key="loc3_save"):
            save_loc_submission("LOC_3", source_text, main_text, reflection)

    elif exercise == "4 Tone & Website/App UX":
        st.subheader("Tone & Website/App UX")
        source_text = "Oops! Something went wrong. Please try again later."
        st.write("Localise this microcopy for a polished Arabic app experience.")
        main_text = st.text_area("Your UX-localised version", key="loc4_main", height=180)
        reflection = st.text_area("Reflection", key="loc4_reflect", height=100)
        if st.button("Save this task", key="loc4_save"):
            save_loc_submission("LOC_4", source_text, main_text, reflection)

    elif exercise == "5 Post-editing: Error Detection":
        st.subheader("Post-editing: Error Detection")
        source_text = "The conference starts on Monday at 9:30 a.m. in Hall B."
        mt_output = "يبدأ المؤتمر يوم الإثنين الساعة 9:30 مساءً في القاعة ب."
        st.write("Source text:")
        st.write(source_text)
        st.write("MT output:")
        st.write(mt_output)
        st.write("Correct the MT output.")
        main_text = st.text_area("Your post-edited version", value=mt_output, key="loc5_main", height=180)
        reflection = st.text_area("Reflection", key="loc5_reflect", height=100)
        if st.button("Save this task", key="loc5_save"):
            save_loc_submission("LOC_5", source_text, main_text, reflection)

    elif exercise == "6 App Store Description":
        st.subheader("App Store Description")
        source_text = (
            "Track your habits, build routines, and stay motivated with daily reminders "
            "and beautifully designed progress charts."
        )
        st.write("Produce a natural Arabic app-store style description.")
        main_text = st.text_area("Your localised app description", key="loc6_main", height=180)
        reflection = st.text_area("Reflection", key="loc6_reflect", height=100)
        if st.button("Save this task", key="loc6_save"):
            save_loc_submission("LOC_6", source_text, main_text, reflection)

    elif exercise == "7 Strategy & Theory Reflection":
        st.subheader("Strategy & Theory Reflection")
        source_text = "Reflect on the difference between translation and localisation in 1–2 short paragraphs."
        st.write(source_text)
        main_text = st.text_area("Your reflection", key="loc7_main", height=220)
        reflection = st.text_area("Optional note on your strategy", key="loc7_reflect", height=100)
        if st.button("Save this task", key="loc7_save"):
            save_loc_submission("LOC_7", source_text, main_text, reflection)

    else:
        st.subheader("Sticker / text / image task")
        loc_stickers = load_json(LOC_STICKERS_FILE)
        if not isinstance(loc_stickers, dict):
            loc_stickers = {}

        if not loc_stickers:
            st.info("No sticker/text/image task has been published by the instructor yet.")
            return

        task_ids = sorted(loc_stickers.keys())
        chosen_task = st.selectbox("Choose instructor task", task_ids, key="sticker_task_select")
        task = loc_stickers[chosen_task]

        st.markdown(f"**Title:** {task.get('title', '')}")
        if task.get("instructions"):
            st.markdown("**Instructions:**")
            st.write(task.get("instructions"))

        source_text = task.get("content_text", "")

        if source_text:
            st.markdown("**Text to localise:**")
            st.write(source_text)

        if task.get("image_type") == "uploaded":
            img_path = task.get("image_path", "")
            if img_path and Path(img_path).exists():
                st.image(img_path)
        elif task.get("image_type") == "url":
            st.image(task.get("image_url", ""))

        main_text = st.text_area("Your localisation / analysis", key="loc_sticker_main", height=220)
        reflection = st.text_area("Reflection", key="loc_sticker_reflect", height=100)

        if st.button("Save this task", key="loc_sticker_save"):
            save_loc_submission(f"LOC_STICKER_{chosen_task}", source_text or task.get("title", ""), main_text, reflection)


# ---------------- Main ----------------
def main():
    st.set_page_config(page_title="Translation Lab (EduApp)", layout="wide")
    st.sidebar.title("Navigation")
    st.sidebar.info(
        f"Loaded: {THIS_FILE}\n\nLast modified: {LAST_EDIT:%Y-%m-%d %H:%M:%S}"
    )

    st.markdown(
        "<div style='padding:8px;border:1px solid #ddd;border-radius:8px;background:#f7f9ff'>"
        "<b>EduApp – Build:</b> 2025-11-10 v7 (translation + localisation + MQM analytics + AI MQM suggestions)</div>",
        unsafe_allow_html=True,
    )

    section = st.sidebar.radio(
        "Module",
        ["Core Translation Lab", "Localisation Lab"],
        index=0,
    )

    if section == "Localisation Lab":
        localisation_lab()
        return

    role = st.sidebar.radio("Login as", ["Instructor", "Student"], index=1)
    if role == "Instructor":
        instructor_dashboard()
    else:
        student_dashboard()


if __name__ == "__main__":
    main()
