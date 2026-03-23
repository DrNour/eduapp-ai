import os
import re
import json
import time
import hashlib
import random
import threading
from io import BytesIO
from pathlib import Path
from typing import List, Tuple
from difflib import SequenceMatcher, ndiff
import datetime

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor

# Optional metrics deps (graceful fallback if missing)
try:
    import sacrebleu
except Exception:
    sacrebleu = None

try:
    from bert_score import score as bertscore_score
except Exception:
    bertscore_score = None

# Optional plotting
try:
    import matplotlib.pyplot as plt  # noqa: F401
    _HAVE_MPL = True
except Exception:
    _HAVE_MPL = False

# Optional OpenAI
try:
    import openai
except Exception:
    openai = None

# ---------------- Proof-of-life banner ----------------
try:
    THIS_FILE = os.path.abspath(__file__)
    LAST_EDIT = datetime.datetime.fromtimestamp(os.path.getmtime(THIS_FILE))
except Exception:
    THIS_FILE = "interactive_session"
    LAST_EDIT = datetime.datetime.now()

# ---------------- Storage ----------------
DATA_DIR = Path("./data")
DATA_DIR.mkdir(exist_ok=True)

EXERCISES_FILE = DATA_DIR / "exercises.json"
SUBMISSIONS_FILE = DATA_DIR / "submissions.json"
LEADERBOARD_FILE = DATA_DIR / "leaderboard.json"

LOC_STICKERS_FILE = DATA_DIR / "loc_stickers.json"
STICKER_IMG_DIR = DATA_DIR / "stickers"
STICKER_IMG_DIR.mkdir(exist_ok=True)

_lock = threading.Lock()


def load_json(file: Path):
    file = Path(file)
    if file.exists():
        with file.open("r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                return {}
    return {}


def save_json(file: Path, data):
    with _lock:
        tmp = Path(str(file) + ".tmp")
        with tmp.open("w", encoding="utf-8") as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
        tmp.replace(file)


# ---------------- Secrets / Auth ----------------
def get_secret(name, default=""):
    value = os.getenv(name, "")
    if value:
        return value
    try:
        if hasattr(st, "secrets") and name in st.secrets:
            return st.secrets[name]
    except Exception:
        pass
    return default


_FALLBACK_PLAIN = "admin123"


def check_password(typed: str) -> bool:
    try:
        instructor_plain = get_secret("INSTRUCTOR_PASSWORD_PLAIN", "")
        instructor_sha256 = get_secret("INSTRUCTOR_PASSWORD_SHA256", "")
        instructor_dev_mode = get_secret("INSTRUCTOR_DEV_MODE", "0") == "1"

        if instructor_sha256:
            h = hashlib.sha256(typed.encode("utf-8")).hexdigest()
            return h == instructor_sha256

        if instructor_plain:
            return typed == instructor_plain

        if instructor_dev_mode:
            return typed == _FALLBACK_PLAIN

        return False
    except Exception:
        return False


def has_instructor_auth_config() -> bool:
    return bool(
        get_secret("INSTRUCTOR_PASSWORD_PLAIN", "")
        or get_secret("INSTRUCTOR_PASSWORD_SHA256", "")
        or get_secret("INSTRUCTOR_DEV_MODE", "0") == "1"
    )


# ---------------- AI backend diagnostics ----------------
def get_ai_backend_status():
    openai_key = get_secret("OPENAI_API_KEY", "")
    openai_available = bool(openai_key)
    openai_model = get_secret("OPENAI_MODEL", "gpt-4.1-mini")

    hf_token = get_secret("HF_API_TOKEN", "")
    hf_available = bool(hf_token)

    return {
        "openai": openai_available,
        "openai_model": openai_model if openai_available else None,
        "hf": hf_available,
    }


# ---------------- Tokenization & Edit Helpers ----------------
_token_re = re.compile(r"\w+|[^\w\s]", re.UNICODE)


def _tokenize(s: str) -> List[str]:
    return _token_re.findall(s or "")


def compute_edit_details(mt_text: str, student_text: str) -> Tuple[int, int, int]:
    mt_tokens = _tokenize(mt_text)
    st_tokens = _tokenize(student_text)
    matcher = SequenceMatcher(None, mt_tokens, st_tokens)

    additions = deletions = replacements = 0
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "insert":
            additions += (j2 - j1)
        elif tag == "delete":
            deletions += (i2 - i1)
        elif tag == "replace":
            replacements += max(i2 - i1, j2 - j1)
    total_edits = additions + deletions + replacements
    return additions, deletions, total_edits


# ---------------- Metrics ----------------
def evaluate_translation(student_text, mt_text=None, reference=None, task_type="Translate", source_text=""):
    src_len = max(1, len(_tokenize(source_text)))
    tgt_len = len(_tokenize(student_text))
    length_ratio = round(tgt_len / src_len, 3)

    if task_type == "Post-edit MT" and mt_text:
        additions, deletions, edits = compute_edit_details(mt_text, student_text)
    else:
        additions = deletions = edits = 0

    bleu = chrf = bert_f1 = None
    if reference:
        refs = [reference]
        try:
            if sacrebleu:
                bleu = float(sacrebleu.corpus_bleu([student_text], [refs]).score)
                chrf = float(sacrebleu.corpus_chrf([student_text], [refs]).score)
        except Exception:
            bleu = None
            chrf = None
        try:
            if bertscore_score:
                P, R, F1 = bertscore_score([student_text], [reference], lang="en")
                bert_f1 = float(F1.mean().item())
        except Exception:
            bert_f1 = None

    return {
        "length_ratio": length_ratio,
        "BLEU": None if bleu is None else round(bleu, 2),
        "chrF++": None if chrf is None else round(chrf, 2),
        "BERTScore_F1": None if bert_f1 is None else round(bert_f1, 3),
        "additions": additions,
        "deletions": deletions,
        "edits": edits
    }


# ---------------- Track Changes ----------------
def _join_tokens_for_display(tokens: List[str]) -> str:
    out = " ".join(tokens)
    out = re.sub(r"\s+([.,!?;:])", r"\1", out)
    return out


def diff_text(baseline: str, student_text: str) -> str:
    differ = ndiff(_tokenize(baseline), _tokenize(student_text))
    parts = []
    for w in differ:
        token = w[2:]
        if w.startswith("- "):
            parts.append(f"<span style='color:#c00;text-decoration:line-through'>{token}</span>")
        elif w.startswith("+ "):
            parts.append(f"<span style='color:#080'>{token}</span>")
        else:
            parts.append(token)
    return _join_tokens_for_display(parts)


_INVALID_XML_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def _safe_docx_text(value) -> str:
    if value is None:
        return ""
    if not isinstance(value, str):
        value = str(value)
    return _INVALID_XML_RE.sub("", value)


def add_diff_to_doc(doc: Document, baseline: str, student_text: str):
    differ = ndiff(_tokenize(baseline), _tokenize(student_text))
    p = doc.add_paragraph()
    for w in differ:
        token = _safe_docx_text(w[2:])
        if w.startswith("- "):
            run = p.add_run(token + " ")
            run.font.strike = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif w.startswith("+ "):
            run = p.add_run(token + " ")
            run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            p.add_run(token + " ")


# ---------------- Exports ----------------
def export_student_word(submissions, student_name):
    doc = Document()
    doc.add_heading(_safe_docx_text(f"Student: {student_name}"), 0)
    subs = submissions.get(student_name, {})
    for ex_id, sub in subs.items():
        doc.add_heading(_safe_docx_text(f"Exercise {ex_id}"), level=1)
        doc.add_paragraph("Source Text:")
        doc.add_paragraph(_safe_docx_text(sub.get("source_text", "")))
        if sub.get("mt_text"):
            doc.add_paragraph("MT Output:")
            doc.add_paragraph(_safe_docx_text(sub.get("mt_text", "")))

        if sub.get("task_type") == "Post-edit MT":
            doc.add_paragraph("Student Submission (Track Changes):")
            base = sub.get("mt_text", "") or ""
            add_diff_to_doc(doc, base, sub.get("student_text", ""))
        else:
            doc.add_paragraph("Student Submission:")
            doc.add_paragraph(_safe_docx_text(sub.get("student_text", "")))

        metrics = sub.get("metrics", {})
        doc.add_paragraph(_safe_docx_text(f"Metrics: {metrics}"))
        doc.add_paragraph(_safe_docx_text(f"Task Type: {sub.get('task_type', '')}"))
        doc.add_paragraph(_safe_docx_text(f"Time Spent: {sub.get('time_spent_sec', 0):.2f} sec"))
        doc.add_paragraph(_safe_docx_text(f"Characters (not keystrokes): {sub.get('keystrokes', 0)}"))
        if sub.get("reflection"):
            doc.add_paragraph("Reflection:")
            doc.add_paragraph(_safe_docx_text(sub.get("reflection")))
        doc.add_paragraph("---")

    buf = BytesIO()
    try:
        doc.save(buf)
    except Exception:
        raise RuntimeError(
            "Failed to generate Word file. Some submission text may contain unsupported characters."
        )
    buf.seek(0)
    return buf


def export_summary_excel(submissions):
    rows = []
    for student, subs in submissions.items():
        for ex_id, sub in subs.items():
            m = sub.get("metrics", {})
            rows.append({
                "Student": student,
                "Exercise": ex_id,
                "Task Type": sub.get("task_type", ""),
                "Length Ratio": m.get("length_ratio"),
                "BLEU": m.get("BLEU"),
                "chrF++": m.get("chrF++"),
                "BERTScore_F1": m.get("BERTScore_F1"),
                "Additions": m.get("additions"),
                "Deletions": m.get("deletions"),
                "Edits": m.get("edits"),
                "Time Spent (s)": sub.get("time_spent_sec", 0),
                "Characters Typed": sub.get("keystrokes", 0)
            })
    df = pd.DataFrame(rows)
    buf = BytesIO()
    df.to_excel(buf, index=False)
    buf.seek(0)
    return buf


# ---------------- Gamification ----------------
def load_leaderboard():
    return load_json(LEADERBOARD_FILE)


def update_leaderboard(student_name, points):
    leaderboard = load_leaderboard()
    leaderboard[student_name] = leaderboard.get(student_name, 0) + points
    save_json(LEADERBOARD_FILE, leaderboard)


def show_leaderboard():
    leaderboard = load_leaderboard()
    st.subheader("Leaderboard")
    if leaderboard:
        items = sorted(leaderboard.items(), key=lambda x: x[1], reverse=True)
        df = pd.DataFrame(items, columns=["Student", "Points"])
        st.dataframe(df, use_container_width=True)
    else:
        st.info("No leaderboard data yet.")


# ---------------- Instructor sticker manager ----------------
def localisation_sticker_manager():
    st.subheader("Instructor – Sticker / text / image localisation tasks")

    pwd = st.text_input("Instructor password", type="password", key="loc_sticker_pwd")
    if not check_password(pwd):
        if not has_instructor_auth_config():
            st.warning(
                "Instructor password is not configured. Set INSTRUCTOR_PASSWORD_PLAIN or "
                "INSTRUCTOR_PASSWORD_SHA256 (or INSTRUCTOR_DEV_MODE=1 for local testing)."
            )
        else:
            st.info("Enter instructor password to manage these tasks.")
        return

    loc_stickers = load_json(LOC_STICKERS_FILE)
    sticker_ids = ["New task"] + sorted(loc_stickers.keys())
    selection = st.selectbox("Choose task", sticker_ids, key="loc_sticker_select")

    if selection != "New task" and selection in loc_stickers:
        current = loc_stickers[selection]
        default_title = current.get("title", "")
        default_instr = current.get("instructions", "")
        default_text = current.get("content_text", "")
        default_url = current.get("image_url", "") if current.get("image_type") == "url" else ""

        st.markdown("**Current preview for students:**")
        if default_text:
            st.markdown("**Text to localise:**")
            st.write(default_text)
        if current.get("image_type") == "uploaded":
            img_path = current.get("image_path", "")
            if img_path and Path(img_path).exists():
                st.image(str(img_path))
            else:
                st.warning("Image file not found on server.")
        elif current.get("image_type") == "url":
            st.image(current.get("image_url", ""))
    else:
        default_title = ""
        default_instr = ""
        default_text = ""
        default_url = ""

    with st.form("loc_sticker_form"):
        title = st.text_input("Task title", value=default_title)
        content_text = st.text_area("Text to be localised (optional)", value=default_text, height=120)
        instructions = st.text_area(
            "Instructions for students (what to do with this text / image)",
            value=default_instr,
            height=120
        )

        st.write("Sticker / image (choose either URL or upload, both optional):")
        image_url = st.text_input("Image URL (optional)", value=default_url)
        uploaded = st.file_uploader("Or upload an image file", type=["png", "jpg", "jpeg", "webp"])

        col1, col2 = st.columns(2)
        with col1:
            save_btn = st.form_submit_button("Save / Update task")
        with col2:
            delete_btn = st.form_submit_button("Delete this task")

    loc_stickers = load_json(LOC_STICKERS_FILE)

    if save_btn:
        if not title.strip() and not content_text.strip() and not image_url and not uploaded:
            st.warning("Please provide at least a title, some text, or an image before saving.")
            return

        if selection != "New task" and selection in loc_stickers:
            task_id = selection
        else:
            existing_nums = []
            for sid in loc_stickers.keys():
                m = re.match(r"STK_(\d+)$", sid)
                if m:
                    existing_nums.append(int(m.group(1)))
            next_num = max(existing_nums + [0]) + 1
            task_id = f"STK_{next_num:03d}"

        image_type = None
        image_path = ""
        image_url_final = ""

        if uploaded is not None:
            safe_name = re.sub(r"[^\w\.\-]", "_", uploaded.name)
            file_path = STICKER_IMG_DIR / f"{task_id}_{safe_name}"
            with open(file_path, "wb") as f:
                f.write(uploaded.getbuffer())
            image_type = "uploaded"
            image_path = str(file_path)
        elif image_url:
            image_type = "url"
            image_url_final = image_url

        if task_id in loc_stickers and image_type is None:
            image_type = loc_stickers[task_id].get("image_type")
            image_path = loc_stickers[task_id].get("image_path", "")
            image_url_final = loc_stickers[task_id].get("image_url", "")

        loc_stickers[task_id] = {
            "title": title.strip(),
            "instructions": instructions.strip(),
            "content_text": content_text.strip(),
            "image_type": image_type,
            "image_path": image_path,
            "image_url": image_url_final,
            "created_at": datetime.datetime.now().isoformat()
        }
        save_json(LOC_STICKERS_FILE, loc_stickers)
        st.success(f"Task {task_id} saved.")

    if delete_btn and selection != "New task" and selection in loc_stickers:
        loc_stickers.pop(selection, None)
        save_json(LOC_STICKERS_FILE, loc_stickers)
        st.success(f"Task {selection} deleted.")


# ---------------- Optional AI exercise generator ----------------
def ai_generate_text(prompt):
    HF_TOKEN = get_secret("HF_API_TOKEN", "")
    if not HF_TOKEN:
        return None

    try:
        import requests
        headers = {"Authorization": f"Bearer {HF_TOKEN}"}
        payload = {"inputs": prompt}
        response = requests.post(
            "https://api-inference.huggingface.co/models/gpt2",
            headers=headers,
            json=payload,
            timeout=15
        )
        if response.status_code == 200:
            data = response.json()
            if isinstance(data, list) and data and "generated_text" in data[0]:
                return data[0]["generated_text"]
    except Exception:
        pass
    return None


# ---------------- Evidence-based Linguistic Hints ----------------
_AR_LETTERS = r"\u0600-\u06FF"


def _tokenize_words(text: str):
    return re.findall(
        r"[A-Za-z" + _AR_LETTERS + r"]+[’'\-]?[A-Za-z" + _AR_LETTERS + r"]+|\d+(?:[.,]\d+)?",
        text
    )


def _likely_terms(source_text: str):
    terms = set()
    for q in re.findall(r"[\"“”‘’'`«»](.+?)[\"“”‘’'`«»]", source_text):
        for w in _tokenize_words(q):
            if len(w) >= 3:
                terms.add(w)

    for w in _tokenize_words(source_text):
        if re.match(r"[A-Z][A-Za-z\-]+$", w):
            terms.add(w)
        elif re.match(r"[A-Z0-9\-]{3,}$", w):
            terms.add(w)
        elif "-" in w or re.search(r"\d", w):
            terms.add(w)
        elif re.match(r"[" + _AR_LETTERS + r"]{4,}$", w):
            terms.add(w)
    return terms


def _short_list(items, n=4):
    items = list(items)
    if not items:
        return ""
    if len(items) <= n:
        return " | ".join(items)
    return " | ".join(items[:n]) + f" … (+{len(items)-n} more)"


def quick_linguistic_hints(source_text: str, student_text: str):
    hints = []
    try:
        src_nums = set(re.findall(r"\d+(?:[.,]\d+)?", source_text))
        tgt_nums = set(re.findall(r"\d+(?:[.,]\d+)?", student_text))
        missing_nums = sorted(src_nums - tgt_nums, key=lambda x: (len(x), x))
        if missing_nums:
            hints.append({
                "rule": "numbers_missing",
                "message": "Some figures from the source didn’t appear in your text.",
                "evidence": f"Missing: {_short_list(missing_nums)}"
            })

        for sym_open, sym_close, label in [
            ("(", ")", "parentheses"),
            ("[", "]", "brackets"),
            ("{", "}", "braces")
        ]:
            if source_text.count(sym_open) != student_text.count(sym_open) or \
               source_text.count(sym_close) != student_text.count(sym_close):
                hints.append({
                    "rule": f"{label}_unbalanced",
                    "message": f"{label.capitalize()} look unbalanced.",
                    "evidence": (
                        f"Source {sym_open}/{sym_close}: "
                        f"{source_text.count(sym_open)}/{source_text.count(sym_close)}; "
                        f"Your text: {student_text.count(sym_open)}/{student_text.count(sym_close)}"
                    )
                })
        if source_text.count('"') != student_text.count('"'):
            hints.append({
                "rule": "quotes_unbalanced",
                "message": "Quotation marks may be unbalanced.",
                "evidence": f"Source quotes: {source_text.count(chr(34))}; Yours: {student_text.count(chr(34))}"
            })

        src_terms = _likely_terms(source_text)
        tgt_tokens = set(_tokenize_words(student_text))
        missing_terms = sorted([t for t in src_terms if t not in tgt_tokens], key=lambda x: (-len(x), x))
        if missing_terms:
            hints.append({
                "rule": "terms_missing",
                "message": "Some key terms/names from the source weren’t reflected.",
                "evidence": f"Examples: {_short_list(missing_terms)}"
            })
    except Exception:
        pass
    return hints


# ---------------- Adaptive Feedback ----------------
def generate_feedback(metrics: dict, task_type: str, source_text: str, student_text: str, extra_hints=None):
    msgs = []
    lr = metrics.get("length_ratio")
    edits = int(metrics.get("edits", 0) or 0)
    adds = int(metrics.get("additions", 0) or 0)
    dels = int(metrics.get("deletions", 0) or 0)
    bleu = metrics.get("BLEU")
    chrf = metrics.get("chrF++")

    if task_type == "Post-edit MT":
        if edits == 0:
            msgs.append((
                "edits_none",
                "No edits were applied to the MT output.",
                "Review the MT carefully—critical errors may remain."
            ))
        elif edits > 20:
            msgs.append((
                "edits_many",
                f"High edit volume detected: {edits} edits (additions {adds}, deletions {dels}).",
                "Prioritize adequacy/accuracy first; avoid cosmetic rephrasing that doesn’t fix meaning."
            ))

    if lr is not None:
        if lr < 0.80:
            msgs.append((
                "len_low",
                f"Length ratio is {lr:.2f} (target ~0.90–1.20).",
                "Your translation may be over-compressed—recheck for omitted content."
            ))
        elif lr > 1.30:
            msgs.append((
                "len_high",
                f"Length ratio is {lr:.2f} (target ~0.90–1.20).",
                "Consider concision—trim redundancy and literal padding."
            ))

    if bleu is not None and chrf is not None:
        if bleu < 30 <= chrf:
            msgs.append((
                "acc_low_flu_ok",
                f"chrF++ is {chrf:.1f} (fluency ok) but BLEU is {bleu:.1f} (accuracy lagging).",
                "Revisit terminology and key meaning units; cross-check against the source."
            ))
        elif bleu >= 30 and chrf < 50:
            msgs.append((
                "flu_low_acc_ok",
                f"BLEU is {bleu:.1f} (accuracy acceptable) but chrF++ is {chrf:.1f} (fluency weak).",
                "Polish cohesion and flow—simplify long clauses and connectors."
            ))
        elif bleu is not None and bleu < 20:
            msgs.append((
                "both_low",
                f"BLEU is {bleu:.1f}.",
                "Start with adequacy: ensure all propositions are conveyed before stylistic edits."
            ))

    if extra_hints:
        for h in extra_hints:
            rule = h.get("rule", "hint")
            msg = h.get("message", "")
            evd = h.get("evidence", "")
            msgs.append((rule, msg, evd))

    seen = set()
    final = []
    for key, text, detail in msgs:
        if key in seen:
            continue
        seen.add(key)
        if detail:
            final.append(f"• {text} — *{detail}*")
        else:
            final.append(f"• {text}")
        if len(final) >= 4:
            break
    return final


# ---------------- AI ----------------
def build_ai_feedback_prompt(source_text: str, mt_text: str, student_text: str, task_type: str) -> str:
    mt_block = mt_text if mt_text else "(no MT output – direct translation task)"
    return f"""
You are an expert English–Arabic translation trainer specialising in translation and MT post-editing.

TASK TYPE: {task_type}

SOURCE TEXT:
{source_text}

MT OUTPUT (if any):
{mt_block}

STUDENT VERSION:
{student_text}

Give concise feedback suitable for a university translation classroom. Please:
1) Comment on accuracy (meaning transfer).
2) Comment on register and appropriateness for the context.
3) Comment on idiomatic and culturally appropriate choices.
4) Highlight one or two concrete examples where the student could improve.
5) If useful, propose a short improved version of one or two sentences, not the entire text.

You may answer partly in Arabic where it helps the student, but keep the structure clear and concise.
"""


def ask_ai_tutor(system_prompt: str, user_prompt: str):
    openai_key = get_secret("OPENAI_API_KEY", "")
    openai_model = get_secret("OPENAI_MODEL", "gpt-4.1-mini")

    if openai and openai_key:
        try:
            if hasattr(openai, "OpenAI"):
                client = openai.OpenAI(api_key=openai_key)

                try:
                    resp = client.responses.create(
                        model=openai_model,
                        instructions=system_prompt,
                        input=user_prompt,
                        max_output_tokens=700,
                    )
                    text = getattr(resp, "output_text", None)
                    if text:
                        return text.strip()
                except Exception:
                    resp = client.chat.completions.create(
                        model=openai_model,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                        max_tokens=700,
                        temperature=0.4,
                    )
                    text = resp.choices[0].message.content
                    if text:
                        return text.strip()

            else:
                openai.api_key = openai_key
                resp = openai.ChatCompletion.create(
                    model=openai_model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    max_tokens=700,
                    temperature=0.4,
                )
                text = resp.choices[0].message["content"]
                if text:
                    return text.strip()

        except Exception as e:
            return f"OpenAI error: {e}"

    hf_token = get_secret("HF_API_TOKEN", "")
    if hf_token:
        try:
            import requests
            headers = {"Authorization": f"Bearer {hf_token}"}
            payload = {"inputs": system_prompt + "\n\n" + user_prompt}
            response = requests.post(
                "https://api-inference.huggingface.co/models/gpt2",
                headers=headers,
                json=payload,
                timeout=30,
            )
            if response.status_code == 200:
                data = response.json()
                if isinstance(data, list) and data and "generated_text" in data[0]:
                    return data[0]["generated_text"].strip()
            return f"Hugging Face error: status {response.status_code}"
        except Exception as e:
            return f"Hugging Face error: {e}"

    return None


def generate_ai_feedback(prompt: str):
    system_prompt = "You are a helpful, expert translation instructor."
    return ask_ai_tutor(system_prompt, prompt)


# ---------------- Instructor Dashboard ----------------
def instructor_dashboard():
    st.title("Instructor Dashboard")
    password = st.text_input("Enter instructor password", type="password")
    if not check_password(password):
        if not has_instructor_auth_config():
            st.warning(
                "Instructor password is not configured. "
                "Set INSTRUCTOR_PASSWORD_PLAIN or INSTRUCTOR_PASSWORD_SHA256 "
                "(or INSTRUCTOR_DEV_MODE=1 for local testing)."
            )
        else:
            st.warning("Incorrect password. Access denied.")
        return

    st.subheader("🔍 AI Feedback Diagnostics")
    status = get_ai_backend_status()

    openai_badge = "❌ Not configured"
    hf_badge = "❌ Not configured"

    if status["openai"]:
        openai_badge = f"✅ Available (model: `{status['openai_model']}`)"
    if status["hf"]:
        hf_badge = "✅ Available (HF_API_TOKEN set)"

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("**OpenAI (ChatGPT API)**")
        st.write(openai_badge)
    with col_b:
        st.markdown("**Hugging Face Inference API**")
        st.write(hf_badge)

    if not status["openai"] and not status["hf"]:
        st.info(
            "No AI backend is currently configured.\n\n"
            "- To use ChatGPT-based feedback, set `OPENAI_API_KEY` (and optional `OPENAI_MODEL`).\n"
            "- To use a Hugging Face model, set `HF_API_TOKEN`.\n\n"
            "The app will automatically prefer OpenAI if both are set."
        )
    else:
        st.success("AI feedback is at least partially configured. Students will be able to request AI feedback.")

    st.markdown("---")

    exercises = load_json(EXERCISES_FILE)
    submissions = load_json(SUBMISSIONS_FILE)

    st.subheader("Create / Edit / Delete Exercise")
    ex_ids = ["New"] + list(exercises.keys())
    selected_ex = st.selectbox("Select Exercise", ex_ids)

    if selected_ex != "New" and selected_ex in exercises:
        default_source = exercises[selected_ex].get("source_text", "")
        default_mt = exercises[selected_ex].get("mt_text", "") or ""
    else:
        default_source = ""
        default_mt = ""

    with st.form("exercise_form"):
        st_text = st.text_area("Source Text", value=default_source, height=150)
        mt_text = st.text_area("MT Output (optional)", value=default_mt, height=150)
        col1, col2, col3 = st.columns(3)
        with col1:
            save_btn = st.form_submit_button("Save Exercise")
        with col2:
            delete_btn = st.form_submit_button("Delete Exercise")
        with col3:
            gen_btn = st.form_submit_button("Generate AI Exercise")

    if save_btn:
        try:
            next_id = (
                str(max([int(k) for k in exercises.keys()] + [0]) + 1).zfill(3)
                if selected_ex == "New" else selected_ex
            )
        except Exception:
            next_id = "001" if selected_ex == "New" else selected_ex

        exercises[next_id] = {
            "source_text": st_text,
            "mt_text": (mt_text.strip() or None)
        }
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise saved! ID: {next_id}")

    if delete_btn and selected_ex != "New":
        exercises.pop(selected_ex, None)
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise {selected_ex} deleted!")

    if gen_btn:
        prompt = "Write a short culturally rich text for translation students."
        ai_text = ai_generate_text(prompt)
        new_text = ai_text if ai_text else f"This is AI-generated exercise {random.randint(1, 1000)}."
        new_mt = f"MT output for exercise {random.randint(1, 1000)}."
        try:
            next_id = str(max([int(k) for k in exercises.keys()] + [0]) + 1).zfill(3)
        except Exception:
            next_id = "001"
        exercises[next_id] = {"source_text": new_text, "mt_text": new_mt}
        save_json(EXERCISES_FILE, exercises)
        st.success(f"Exercise saved as ID {next_id}")

    st.subheader("Download Exercises")
    if exercises:
        for ex_id, ex in exercises.items():
            try:
                buf = BytesIO()
                doc = Document()
                doc.add_heading(f"Exercise {ex_id}", 0)
                doc.add_paragraph("Source Text:")
                doc.add_paragraph(ex.get("source_text", ""))
                if ex.get("mt_text"):
                    doc.add_paragraph("MT Output:")
                    doc.add_paragraph(ex.get("mt_text", ""))
                doc.save(buf)
                buf.seek(0)
                st.download_button(
                    f"Exercise {ex_id} (Word)",
                    buf,
                    file_name=f"Exercise_{ex_id}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception:
                st.info(f"Exercise {ex_id}: export not available (DOCX error).")
    else:
        st.info("No exercises yet.")

    st.subheader("Student Submissions & Exports")
    if submissions:
        student_choice = st.selectbox("Choose student", ["All"] + list(submissions.keys()))
        if student_choice != "All":
            try:
                buf = export_student_word(submissions, student_choice)
                safe_name = re.sub(r"[^\w\-]+", "_", student_choice)
                st.download_button(
                    f"Download {student_choice}'s Submissions (Word)",
                    buf,
                    file_name=f"{safe_name}_submissions.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except RuntimeError as e:
                st.warning(str(e))
            except Exception:
                st.error("Unexpected error while generating the Word report.")

        st.subheader("Download Metrics Summary (Excel)")
        excel_buf = export_summary_excel(submissions)
        st.download_button(
            "Download Excel Summary",
            excel_buf,
            file_name="metrics_summary.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        try:
            st.subheader("Class Snapshot")
            rows = []
            for ex_id2, ex in exercises.items():
                vals = []
                for student, subs in submissions.items():
                    sub = subs.get(ex_id2)
                    if sub:
                        m = sub.get("metrics", {})
                        if m.get("chrF++") is not None:
                            vals.append(m["chrF++"])
                if vals:
                    mean_val = round(sum(vals) / max(1, len(vals)), 2)
                    rows.append({"Exercise": ex_id2, "chrF++ mean": mean_val, "n": len(vals)})
            if rows:
                st.dataframe(pd.DataFrame(rows))
            else:
                st.info("No metrics yet to summarize.")
        except Exception:
            st.info("Snapshot unavailable (aggregation error).")

        show_leaderboard()
    else:
        st.info("No submissions yet.")


# ---------------- Student Dashboard ----------------
def student_dashboard():
    st.title("Student Dashboard")

    exercises = load_json(EXERCISES_FILE)
    if not exercises:
        st.info("No exercises available yet. Please check back later.")
        return

    submissions = load_json(SUBMISSIONS_FILE)
    student_name = st.text_input("Enter your name")
    if not student_name:
        return

    if student_name not in submissions:
        submissions[student_name] = {}

    ex_id = st.selectbox("Choose Exercise", list(exercises.keys()))
    if not ex_id:
        return

    ex = exercises[ex_id]
    existing_sub = submissions.get(student_name, {}).get(ex_id, {})

    st.subheader("Source Text")
    st.markdown(
        f"<div style='font-family:Times New Roman;font-size:12pt;'>{ex.get('source_text', '')}</div>",
        unsafe_allow_html=True
    )

    task_options = ["Translate"] if not ex.get("mt_text") else ["Translate", "Post-edit MT"]

    default_task_type = existing_sub.get("task_type", task_options[0])
    if default_task_type not in task_options:
        default_task_type = task_options[0]

    task_type = st.radio(
        "Task Type",
        task_options,
        horizontal=True,
        index=task_options.index(default_task_type)
    )

    if existing_sub.get("student_text"):
        initial_text = existing_sub.get("student_text", "")
    else:
        initial_text = "" if task_type == "Translate" else (ex.get("mt_text", "") or "")

    default_reflection = existing_sub.get("reflection", "")

    start_key = f"start_time_{student_name}_{ex_id}"
    keys_key = f"chars_{student_name}_{ex_id}"
    if start_key not in st.session_state:
        st.session_state[start_key] = time.time()
    if keys_key not in st.session_state:
        st.session_state[keys_key] = 0

    with st.form(key=f"exercise_form_{student_name}_{ex_id}"):
        student_text = st.text_area(
            "Type your translation / post-edit here",
            value=initial_text,
            height=300
        )
        reflection = st.text_area(
            "Brief reflection (what changed / why?)",
            value=default_reflection,
            height=80
        )
        submitted = st.form_submit_button("Submit")

    if submitted:
        time_spent = time.time() - st.session_state[start_key]
        st.session_state[keys_key] = len(student_text)

        metrics = evaluate_translation(
            student_text,
            mt_text=ex.get("mt_text"),
            reference=None,
            task_type=task_type,
            source_text=ex.get("source_text", ""),
        )

        submissions[student_name][ex_id] = {
            "source_text": ex.get("source_text", ""),
            "mt_text": ex.get("mt_text"),
            "student_text": student_text,
            "task_type": task_type,
            "time_spent_sec": round(time_spent, 2),
            "keystrokes": st.session_state[keys_key],
            "metrics": metrics,
            "reflection": reflection,
        }
        save_json(SUBMISSIONS_FILE, submissions)

        points = 0
        try:
            if metrics.get("BLEU") is not None:
                points += int(metrics["BLEU"])
            if metrics.get("chrF++") is not None:
                points += int(metrics["chrF++"] / 2)
            if task_type == "Post-edit MT":
                points += max(0, 10 - int(metrics["edits"]))
        except Exception:
            pass
        update_leaderboard(student_name, points)

        st.success("Submission saved!")
        existing_sub = submissions[student_name][ex_id]

    if not existing_sub or not existing_sub.get("student_text", "").strip():
        st.info("Write your translation and click Submit first. Then AI feedback and AI Tutor will work.")
        return

    student_text = existing_sub.get("student_text", "")
    reflection = existing_sub.get("reflection", "")
    task_type = existing_sub.get("task_type", task_type)
    time_spent = existing_sub.get("time_spent_sec", 0)
    st.session_state[keys_key] = existing_sub.get("keystrokes", len(student_text))
    metrics = existing_sub.get("metrics", {})

    def _fmt(v):
        return "—" if v is None else v

    st.subheader("Your Metrics")
    st.markdown(f"""
- **Length Ratio** (target/src): {_fmt(metrics.get('length_ratio'))}
- **BLEU**: {_fmt(metrics.get('BLEU'))}
- **chrF++**: {_fmt(metrics.get('chrF++'))}
- **BERTScore F1**: {_fmt(metrics.get('BERTScore_F1'))}
- **Additions**: {_fmt(metrics.get('additions'))}
- **Deletions**: {_fmt(metrics.get('deletions'))}
- **Edits**: {_fmt(metrics.get('edits'))}
- **Time Spent**: {time_spent} sec
- **Characters Typed**: {st.session_state[keys_key]}
""")

    extra = quick_linguistic_hints(ex.get("source_text", ""), student_text)
    feedback_msgs = generate_feedback(
        metrics,
        task_type,
        ex.get("source_text", ""),
        student_text,
        extra
    )

    st.subheader("Adaptive Feedback (metrics-based)")
    if feedback_msgs:
        for m in feedback_msgs:
            st.markdown(m)
    else:
        st.info("No specific issues triggered. Focus on cohesion, clarity, and consistent terminology.")

    if task_type == "Post-edit MT":
        st.subheader("Track Changes")
        st.caption("Track changes: green = additions, red strike = deletions.")
        base = ex.get("mt_text", "") or ""
        st.markdown(diff_text(base, student_text), unsafe_allow_html=True)

    try:
        history = []
        for ex_id2, sub2 in submissions.get(student_name, {}).items():
            m2 = sub2.get("metrics", {})
            history.append({
                "ex": ex_id2,
                "BLEU": m2.get("BLEU"),
                "chrF++": m2.get("chrF++"),
                "Edits": m2.get("edits", 0)
            })
        if history:
            st.subheader("Progress Overview")
            df_hist = pd.DataFrame(history)
            try:
                if not df_hist.empty:
                    df_trend = df_hist.set_index("ex")[["BLEU", "chrF++"]]
                    st.line_chart(df_trend)
            except Exception:
                pass
            try:
                df_edits = df_hist.set_index("ex")[["Edits"]]
                st.bar_chart(df_edits)
            except Exception:
                pass
    except Exception:
        st.info("Progress charts unavailable.")

    show_leaderboard()

    st.subheader("Optional AI Feedback (experimental)")
    if st.button("Get AI feedback on your submission"):
        prompt = build_ai_feedback_prompt(
            ex.get("source_text", ""),
            ex.get("mt_text", "") or "",
            student_text,
            task_type,
        )
        with st.spinner("Requesting AI feedback..."):
            ai_text = generate_ai_feedback(prompt)

        if ai_text:
            st.markdown("### AI feedback / suggestion")
            st.write(ai_text)
        else:
            st.error("AI call returned no text. Check API key, model name, and OpenAI package version.")

    st.subheader("AI Tutor Chat")
    student_prompt = st.text_area(
        "Ask the AI tutor a question about your translation",
        height=120,
        placeholder="Example: Is my Arabic too literal? Suggest 2 more idiomatic options for the first sentence.",
        key=f"ai_tutor_prompt_{student_name}_{ex_id}",
    )

    if st.button("Ask AI Tutor"):
        if not student_prompt.strip():
            st.warning("Please enter a prompt.")
        else:
            system_prompt = """
You are an expert English-Arabic translation tutor for university students.
Help students improve their work.
Do not automatically complete the whole task unless explicitly requested.
Focus on accuracy, register, idiomaticity, and cultural appropriateness.
Keep responses clear and concise.
You may use both English and Arabic.
"""

            user_prompt = f"""
TASK TYPE: {task_type}

SOURCE TEXT:
{ex.get('source_text', '')}

MT OUTPUT:
{ex.get('mt_text', '') or '(none)'}

STUDENT SUBMISSION:
{student_text}

STUDENT QUESTION:
{student_prompt}
"""

            with st.spinner("Requesting AI response..."):
                ai_text = ask_ai_tutor(system_prompt, user_prompt)

            if ai_text:
                st.markdown("### AI Tutor Response")
                st.write(ai_text)
            else:
                st.error("AI call returned no text. Check API key, model name, and OpenAI package version.")


# ---------------- Localisation Lab ----------------
def localisation_lab():
    st.title("🌍 Localisation Lab")
    st.write(
        "Interactive exercises on localisation (English ↔ Arabic). "
        "Work here is saved to the same JSON/leaderboard as the core lab."
    )

    mode = st.sidebar.radio(
        "Localisation mode",
        ["Student view", "Instructor (manage sticker/text/image tasks)"],
        index=0,
        key="loc_mode"
    )

    if mode == "Instructor (manage sticker/text/image tasks)":
        localisation_sticker_manager()
        return

    student_name = st.text_input("Enter your name (for saving localisation work)")
    if not student_name:
        st.info("Please enter your name to start.")
        return

    submissions = load_json(SUBMISSIONS_FILE)
    if student_name not in submissions:
        submissions[student_name] = {}

    exercise = st.sidebar.selectbox(
        "Choose a localisation exercise",
        [
            "1️⃣ Translation vs Localisation",
            "2️⃣ Cultural Adaptation in Advertising",
            "3️⃣ Conventions: Dates, Units, Currency",
            "4️⃣ Tone & Website/App UX",
            "5️⃣ Post-editing: Error Detection",
            "6️⃣ App Store Description",
            "7️⃣ Strategy & Theory Reflection",
            "🎨 Sticker / text / image task (from instructor)",
        ],
        key="loc_ex_select",
    )

    def save_loc_submission(ex_id: str, source_text: str, main_text: str, reflection_text: str):
        if not main_text.strip():
            st.warning("Nothing to save yet — please write your main answer first.")
            return

        start_key = f"loc_start_{student_name}_{ex_id}"
        if start_key not in st.session_state:
            st.session_state[start_key] = time.time()
        time_spent = time.time() - st.session_state[start_key]

        keystrokes = len(main_text)

        metrics = evaluate_translation(
            main_text,
            mt_text=None,
            reference=None,
            task_type="Localisation",
            source_text=source_text,
        )

        submissions[student_name][ex_id] = {
            "source_text": source_text,
            "mt_text": None,
            "student_text": main_text,
            "task_type": "Localisation",
            "time_spent_sec": round(time_spent, 2),
            "keystrokes": keystrokes,
            "metrics": metrics,
            "reflection": reflection_text,
        }
        save_json(SUBMISSIONS_FILE, submissions)

        points = 15
        lr = metrics.get("length_ratio")
        try:
            if lr is not None and 0.8 <= lr <= 1.3:
                points += 5
        except Exception:
            pass
        update_leaderboard(student_name, points)

        st.success("Localisation submission saved and leaderboard updated!")

        def _fmt(v):
            return "—" if v is None else v

        st.subheader("Your Metrics (Localisation)")
        st.write(f"• Length Ratio (target/src): {_fmt(metrics['length_ratio'])}")
        st.write(f"• BLEU: {_fmt(metrics['BLEU'])}")
        st.write(f"• chrF++: {_fmt(metrics['chrF++'])}")
        st.write(f"• BERTScore F1: {_fmt(metrics['BERTScore_F1'])}")
        st.write(f"• Time Spent: {round(time_spent, 2)} sec")
        st.write(f"• Characters Typed: {keystrokes}")

        extra = quick_linguistic_hints(source_text, main_text)
        feedback_msgs = generate_feedback(
            metrics,
            "Localisation",
            source_text,
            main_text,
            extra_hints=extra,
        )

        st.subheader("Adaptive Feedback")
        if feedback_msgs:
            for m in feedback_msgs:
                st.markdown(m)
        else:
            st.info("No specific issues triggered. Focus on cohesion, clarity, and consistent localisation choices.")

        st.subheader("Leaderboard (including localisation tasks)")
        show_leaderboard()

        st.subheader("Optional AI Feedback (Localisation)")
        if st.button("Get AI feedback on this localisation task", key=f"ai_loc_{ex_id}"):
            prompt = build_ai_feedback_prompt(source_text, "", main_text, "Localisation")
            with st.spinner("Requesting AI feedback..."):
                ai_text = generate_ai_feedback(prompt)
            if ai_text:
                st.markdown("### AI feedback / suggestion (Localisation)")
                st.write(ai_text)
            else:
                st.warning(
                    "AI feedback is not configured or temporarily unavailable.\n\n"
                    "To enable it, set OPENAI_API_KEY (for ChatGPT) or HF_API_TOKEN (for Hugging Face)."
                )

    ex_id_map = {
        "1️⃣ Translation vs Localisation": "LOC_1",
        "2️⃣ Cultural Adaptation in Advertising": "LOC_2",
        "3️⃣ Conventions: Dates, Units, Currency": "LOC_3",
        "4️⃣ Tone & Website/App UX": "LOC_4",
        "5️⃣ Post-editing: Error Detection": "LOC_5",
        "6️⃣ App Store Description": "LOC_6",
        "7️⃣ Strategy & Theory Reflection": "LOC_7",
        "🎨 Sticker / text / image task (from instructor)": "LOC_STICKER",
    }
    current_ex_id = ex_id_map.get(exercise, "LOC_STICKER")
    start_key = f"loc_start_{student_name}_{current_ex_id}"
    if start_key not in st.session_state:
        st.session_state[start_key] = time.time()

    st.info("Localisation exercise functions omitted here for brevity; plug in your existing ones.")


# ---------------- Main ----------------
def main():
    st.set_page_config(page_title="Translation Lab (EduApp)", layout="wide")
    st.sidebar.title("Navigation")
    st.sidebar.info(
        f"Loaded: {THIS_FILE}\n\nLast modified: {LAST_EDIT:%Y-%m-%d %H:%M:%S}"
    )

    st.markdown(
        "<div style='padding:8px;border:1px solid #ddd;border-radius:8px;background:#f7f9ff'>"
        "<b>EduApp – Build:</b> 2025-11-10 v4 (translation + localisation lab + AI feedback)</div>",
        unsafe_allow_html=True
    )

    section = st.sidebar.radio(
        "Module",
        ["Core Translation Lab", "Localisation Lab"],
        index=0
    )

    if section == "Localisation Lab":
        localisation_lab()
        return

    role = st.sidebar.radio("Login as", ["Instructor", "Student"], index=1)
    if role == "Instructor":
        instructor_dashboard()
    else:
        student_dashboard()


if __name__ == "__main__":
    main()
